from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from datetime import datetime, timedelta
import os
import zipfile
import json
from docx.shared import Inches, Pt, RGBColor
import pdfkit
from io import BytesIO
import uuid
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import io
import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils.dataframe import dataframe_to_rows
import csv
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import qn
from PIL import Image, ImageDraw, ImageFont
import textwrap
import base64
from docx.enum.table import WD_TABLE_ALIGNMENT
import tempfile
import requests

# Инициализация приложения
app = Flask(__name__)
app.config['SECRET_KEY'] = 'olympiad-system-secret-key'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///olympiad.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = 'static/pdf_files'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max upload

if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

# Инициализация расширений
db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'

import json


# Функция для получения текущего локального времени
def get_current_time():
    """Возвращает текущее локальное время"""
    return datetime.now()


@app.template_filter('fromjson')
def fromjson(value):
    return json.loads(value)


# Фильтры для шаблонов
@app.template_filter('tojson')
def to_json(value):
    return json.dumps(value)


def prepare_question_data(questions):
    """Обновленная функция для корректной обработки JSON-полей перед отправкой в шаблон"""
    for q in questions:
        if q.question_type == 'test' and q.options:
            q.options_list = json.loads(q.options)
            if q.correct_answers:
                q.correct_answers_list = json.loads(q.correct_answers)
            else:
                q.correct_answers_list = []
        elif q.question_type == 'matching' and q.matches:
            matches_data = json.loads(q.matches)

            # Проверяем новый или старый формат
            if isinstance(matches_data, dict) and 'left_items' in matches_data:
                # Новый формат
                q.matches_data = matches_data
                q.matches_list = []  # Для обратной совместимости с шаблонами
            else:
                # Старый формат - конвертируем
                q.matches_list = matches_data
                q.matches_data = {
                    'left_items': [match['left'] for match in matches_data],
                    'right_items': [match['right'] for match in matches_data],
                    'correct_matches': {match['left']: match['right'] for match in matches_data}
                }
        else:
            q.options_list = []
            q.matches_list = []
            q.matches_data = {}
            q.correct_answers_list = []
    return questions


def resize_signature_with_aspect_ratio(img, max_width, max_height):
    """
    Изменяет размер изображения подписи с сохранением пропорций

    Args:
        img: PIL Image объект
        max_width: максимальная ширина
        max_height: максимальная высота

    Returns:
        PIL Image: изображение с новым размером
    """
    original_width, original_height = img.size

    # Вычисляем коэффициенты масштабирования для ширины и высоты
    width_ratio = max_width / original_width
    height_ratio = max_height / original_height

    # Используем меньший коэффициент для сохранения пропорций
    scale_ratio = min(width_ratio, height_ratio)

    # Вычисляем новые размеры
    new_width = int(original_width * scale_ratio)
    new_height = int(original_height * scale_ratio)

    # Изменяем размер с высоким качеством
    return img.resize((new_width, new_height), Image.Resampling.LANCZOS)


def get_font(size, bold=False):
    """Получает шрифт нужного размера"""
    try:
        if bold:
            return ImageFont.truetype("arial.ttf", size)
        else:
            return ImageFont.truetype("arial.ttf", size)
    except:
        try:
            if bold:
                return ImageFont.truetype("/System/Library/Fonts/Arial Bold.ttf", size)
            else:
                return ImageFont.truetype("/System/Library/Fonts/Arial.ttf", size)
        except:
            return ImageFont.load_default()


def add_signatures_to_certificate(img, signatures_folder='static/signatures'):
    """Добавляет подписи членов жюри на сертификат"""
    draw = ImageDraw.Draw(img)
    width, height = img.size

    # Позиции для подписей (внизу сертификата)
    signature_y = height - 400
    signature_width = 350
    signature_height = 150

    # Данные членов жюри
    jury_members = [
        {"name": "Мохнатко Ирина Николаевна", "position": "к.т.н., доцент, зав. кафедрой «Гражданская безопасность»",
         "file": "1.jpg"},
        {"name": "Малюта Сергей Иванович", "position": "к.т.н., доцент кафедры «Гражданская безопасность»",
         "file": "2.jpg"},
        {"name": "Мазилин Сергей Дмитриевич", "position": "к.т.н., доцент кафедры «Гражданская безопасность»",
         "file": "3.jpg"}
    ]

    # Расчет позиций для размещения подписей
    spacing_between_signatures = 200
    total_width = len(jury_members) * signature_width + (len(jury_members) - 1) * spacing_between_signatures
    start_x = (width - total_width) // 2

    font_name = get_font(32, bold=True)
    font_position = get_font(24)

    for i, member in enumerate(jury_members):
        x = start_x + i * (signature_width + spacing_between_signatures)

        # Пытаемся загрузить изображение подписи
        try:
            signature_path = os.path.join(signatures_folder, member["file"])
            if os.path.exists(signature_path):
                signature_img = Image.open(signature_path)

                # Масштабируем подпись с сохранением пропорций
                signature_img = resize_signature_with_aspect_ratio(
                    signature_img,
                    max_width=signature_width - 50,
                    max_height=signature_height - 40
                )

                # Центрируем подпись в выделенной области
                signature_area_width = signature_width - 50
                signature_area_height = signature_height - 80

                # Вычисляем позицию для центрирования
                sig_x = x + 25 + (signature_area_width - signature_img.width) // 2
                sig_y = signature_y - 100 + (signature_area_height - signature_img.height) // 2

                # Вставляем подпись
                img.paste(signature_img, (sig_x, sig_y), signature_img if signature_img.mode == 'RGBA' else None)
        except Exception as e:
            print(f"Не удалось загрузить подпись {member['file']}: {e}")
            # Рисуем прямоугольник для подписи
            draw.rectangle([x + 25, signature_y - 100, x + signature_width - 25, signature_y - 20],
                           outline='#CCCCCC', width=2)
            draw.text((x + signature_width // 2, signature_y - 60), "(подпись)",
                      font=font_position, fill='#666666', anchor="mm")

        # Добавляем линию для подписи
        draw.line([x, signature_y, x + signature_width, signature_y], fill='#000000', width=3)

        # Добавляем имя и должность
        name_lines = textwrap.wrap(member["name"], width=30)
        position_lines = textwrap.wrap(member["position"], width=35)

        current_y = signature_y + 20
        for line in name_lines:
            bbox = draw.textbbox((0, 0), line, font=font_name)
            text_width = bbox[2] - bbox[0]
            draw.text((x + signature_width // 2 - text_width // 2, current_y), line,
                      font=font_name, fill='#000000')
            current_y += 45

        current_y += 15
        for line in position_lines:
            bbox = draw.textbbox((0, 0), line, font=font_position)
            text_width = bbox[2] - bbox[0]
            draw.text((x + signature_width // 2 - text_width // 2, current_y), line,
                      font=font_position, fill='#000000')
            current_y += 35

    return img


def generate_certificate(user_name, olympiad_title, olympiad_end_date, certificate_type='participation',
                         place=None, score=None, time_bonus=None, speciality=None,
                         background_image_path='static/images/certificate_background.png'):
    """
    Универсальная функция для генерации сертификатов и дипломов

    Args:
        user_name (str): ФИО участника
        olympiad_title (str): Название олимпиады
        olympiad_end_date: Дата окончания олимпиады
        certificate_type (str): Тип документа ('participation' или 'winner')
        place (int): Место (для дипломов)
        score (float): Итоговые баллы участника
        time_bonus (float): Временной бонус/штраф
        speciality (str): Специальность участника
        background_image_path (str): Путь к фоновому изображению

    Returns:
        PIL.Image: Изображение сертификата/диплома
    """

    # Загружаем готовое фоновое изображение
    try:
        img = Image.open(background_image_path)
        print(f"Загружено фоновое изображение: {background_image_path}")
    except Exception as e:
        print(f"Ошибка загрузки фонового изображения {background_image_path}: {e}")
        # Создаем простой белый фон как fallback
        img = Image.new('RGB', (3508, 2480), 'white')

    if certificate_type == 'winner':
        return _generate_winner_content(img, user_name, olympiad_title, olympiad_end_date,
                                        place, score, speciality)
    else:
        return _generate_participation_content(img, user_name, olympiad_title, olympiad_end_date,
                                               speciality, score, time_bonus)


# Обновленная функция генерации сертификата участия с баллами
def _generate_participation_content(img, user_name, olympiad_title, olympiad_end_date, speciality, score=None,
                                    time_bonus=None):
    """Генерирует содержимое сертификата участия с результатами"""
    draw = ImageDraw.Draw(img)
    width, height = img.size

    # Заголовок университета
    university_lines = [
        "ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ",
        "ВЫСШЕГО ОБРАЗОВАНИЯ «МЕЛИТОПОЛЬСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ»",
        "Технический факультет",
        "кафедра «Гражданская безопасность»"
    ]

    font_header = get_font(48, bold=True)
    font_subheader = get_font(40, bold=True)
    font_small_header = get_font(36, bold=True)

    y = 200
    for i, line in enumerate(university_lines):
        if i < 2:
            current_font = font_header
        elif i == 2:
            current_font = font_subheader
        else:
            current_font = font_small_header

        bbox = draw.textbbox((0, 0), line, font=current_font)
        text_width = bbox[2] - bbox[0]
        draw.text((width // 2 - text_width // 2, y), line, font=current_font, fill='#000000')
        y += 70

    # Заголовок сертификата - УВЕЛИЧЕННЫЙ РАЗМЕР И ЖИРНЫЙ
    y += 100
    certificate_title = "СЕРТИФИКАТ УЧАСТНИКА"
    font_title = get_font(100, bold=True)
    bbox = draw.textbbox((0, 0), certificate_title, font=font_title)
    text_width = bbox[2] - bbox[0]
    draw.text((width // 2 - text_width // 2, y), certificate_title, font=font_title, fill='#820000')

    # Основной текст
    y += 200
    font_main = get_font(48)
    font_name = get_font(56, bold=True)

    # "Настоящим подтверждается, что"
    confirm_text = "Настоящим подтверждается, что"
    bbox = draw.textbbox((0, 0), confirm_text, font=font_main)
    text_width = bbox[2] - bbox[0]
    draw.text((width // 2 - text_width // 2, y), confirm_text, font=font_main, fill='#000000')

    # Имя участника - УВЕЛИЧЕННЫЙ РАЗМЕР
    y += 120
    font_name_big = get_font(70, bold=True)
    bbox = draw.textbbox((0, 0), user_name, font=font_name_big)
    text_width = bbox[2] - bbox[0]
    draw.text((width // 2 - text_width // 2, y), user_name, font=font_name_big, fill='#820000')

    # Подчеркивание имени
    line_start = width // 2 - text_width // 2 - 50
    line_end = width // 2 + text_width // 2 + 50
    draw.line([line_start, y + 80, line_end, y + 80], fill='#820000', width=4)

    # Курс студента
    y += 120
    user = User.query.filter_by(full_name=user_name).first()
    if user and user.course:
        course_text = f"студент {user.course} курса"
        bbox = draw.textbbox((0, 0), course_text, font=font_main)
        text_width = bbox[2] - bbox[0]
        draw.text((width // 2 - text_width // 2, y), course_text, font=font_main, fill='#000000')
        y += 80

    # Специальность (если указана)
    if speciality:
        y += 80
        speciality_text = f"направление подготовки: {speciality}"
        speciality_lines = textwrap.wrap(speciality_text, width=60)
        for line in speciality_lines:
            bbox = draw.textbbox((0, 0), line, font=font_main)
            text_width = bbox[2] - bbox[0]
            draw.text((width // 2 - text_width // 2, y), line, font=font_main, fill='#000000')
            y += 60

    # Текст об участии в олимпиаде
    y += 80
    participation_lines = [
        "принял(а) участие в олимпиаде",
        f'"{olympiad_title.upper()}"'  # БОЛЬШИМИ БУКВАМИ
    ]

    for line in participation_lines:
        if line.startswith('"'):
            current_font = get_font(70, bold=True)  # Увеличен размер названия олимпиады
            color = '#820000'
        else:
            current_font = font_main
            color = '#000000'

        bbox = draw.textbbox((0, 0), line, font=current_font)
        text_width = bbox[2] - bbox[0]
        draw.text((width // 2 - text_width // 2, y), line, font=current_font, fill=color)
        y += 100 if line.startswith('"') else 80

    # НОВОЕ: Блок с результатами участника
    if score is not None:
        y += 80

        # Заголовок результатов
        results_title = "и показал следующий результат:"
        bbox = draw.textbbox((0, 0), results_title, font=font_main)
        text_width = bbox[2] - bbox[0]
        draw.text((width // 2 - text_width // 2, y), results_title, font=font_main, fill='#000000')
        y += 100

        # Рамка для результатов
        box_width = 800
        box_height = 120
        box_x = width // 2 - box_width // 2
        box_y = y

        # Рисуем рамку
        draw.rectangle([box_x, box_y, box_x + box_width, box_y + box_height],
                       outline='#820000', width=4, fill='#FFF8DC')

        # Основной результат
        y += 30
        main_score_text = f"{score:.1f} баллов"
        font_score = get_font(60, bold=True)
        bbox = draw.textbbox((0, 0), main_score_text, font=font_score)
        text_width = bbox[2] - bbox[0]
        draw.text((width // 2 - text_width // 2, y), main_score_text, font=font_score, fill='#820000')


        y += 80

    # Дата в правый угол
    y += 150

    # Определяем дату для вывода
    if hasattr(olympiad_end_date, 'strftime'):
        date_to_format = olympiad_end_date
    else:
        date_to_format = datetime.now()

    formatted_date = date_to_format.strftime("«%d» %B %Y г.")

    # Заменяем английские названия месяцев на русские
    months = {
        'January': 'января', 'February': 'февраля', 'March': 'марта',
        'April': 'апреля', 'May': 'мая', 'June': 'июня',
        'July': 'июля', 'August': 'августа', 'September': 'сентября',
        'October': 'октября', 'November': 'ноября', 'December': 'декабря'
    }
    for eng, rus in months.items():
        formatted_date = formatted_date.replace(eng, rus)

    date_font = get_font(36)
    bbox = draw.textbbox((0, 0), formatted_date, font=date_font)
    text_width = bbox[2] - bbox[0]
    date_x = width - text_width - 200
    date_y = y

    draw.text((date_x, date_y), formatted_date, font=date_font, fill='#2F4F4F')

    # Добавляем подписи
    img = add_signatures_to_certificate(img)

    return img


def _generate_winner_content(img, user_name, olympiad_title, olympiad_end_date, place, score, speciality):
    """Генерирует содержимое диплома победителя"""
    draw = ImageDraw.Draw(img)
    width, height = img.size

    # Заголовок университета
    university_lines = [
        "ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ",
        "ВЫСШЕГО ОБРАЗОВАНИЯ «МЕЛИТОПОЛЬСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ»"
    ]

    faculty_lines = [
        "Технический факультет",
        "кафедра «Гражданская безопасность»"
    ]

    font_header = get_font(52, bold=True)
    font_subheader = get_font(44, bold=True)
    font_small_header = get_font(40, bold=True)

    y = 220
    # Рисуем заголовки университета
    for i, line in enumerate(university_lines):
        current_font = font_header
        color = '#2F4F4F'

        bbox = draw.textbbox((0, 0), line, font=current_font)
        text_width = bbox[2] - bbox[0]
        draw.text((width // 2 - text_width // 2, y), line, font=current_font, fill=color)
        y += 65

    # Отступ между университетом и факультетом
    y += 40

    # Рисуем заголовки факультета
    for i, line in enumerate(faculty_lines):
        if i == 0:
            current_font = font_subheader
            color = '#8B0000'
        else:
            current_font = font_small_header
            color = '#8B0000'

        bbox = draw.textbbox((0, 0), line, font=current_font)
        text_width = bbox[2] - bbox[0]
        draw.text((width // 2 - text_width // 2, y), line, font=current_font, fill=color)
        y += 65

    # Заголовок диплома - УВЕЛИЧЕННЫЙ РАЗМЕР И ЖИРНЫЙ
    y += 120
    certificate_title = "ДИПЛОМ ПОБЕДИТЕЛЯ"
    title_color = '#8B0000'

    font_title = get_font(110, bold=True)  # Увеличен с 90 до 110
    bbox = draw.textbbox((0, 0), certificate_title, font=font_title)
    text_width = bbox[2] - bbox[0]

    # Тень для заголовка
    shadow_offset = 4
    draw.text((width // 2 - text_width // 2 + shadow_offset, y + shadow_offset),
              certificate_title, font=font_title, fill='#CCCCCC')
    draw.text((width // 2 - text_width // 2, y), certificate_title, font=font_title, fill=title_color)

    # Основной текст
    y += 220
    font_main = get_font(52)
    font_name = get_font(60, bold=True)
    font_emphasis = get_font(56, bold=True)

    # "Награждается"
    award_text = "Награждается"
    bbox = draw.textbbox((0, 0), award_text, font=font_main)
    text_width = bbox[2] - bbox[0]
    draw.text((width // 2 - text_width // 2, y), award_text, font=font_main, fill='#2F4F4F')

    # Имя участника - УВЕЛИЧЕННЫЙ РАЗМЕР
    y += 120
    font_name_big = get_font(75, bold=True)  # Увеличен с 60 до 75
    bbox = draw.textbbox((0, 0), user_name, font=font_name_big)
    text_width = bbox[2] - bbox[0]
    draw.text((width // 2 - text_width // 2, y), user_name, font=font_name_big, fill='#8B0000')

    # Курс студента - НОВОЕ
    y += 120
    user = User.query.filter_by(full_name=user_name).first()
    if user and user.course:
        student_text = f"студент {user.course} курса"
    else:
        student_text = "студент 1 курса"  # fallback

    bbox = draw.textbbox((0, 0), student_text, font=font_main)
    text_width = bbox[2] - bbox[0]
    draw.text((width // 2 - text_width // 2, y), student_text, font=font_main, fill='#2F4F4F')

    # Специальность
    if speciality:
        y += 90
        speciality_text = f"направление подготовки: {speciality}"
        speciality_lines = textwrap.wrap(speciality_text, width=55)
        for line in speciality_lines:
            bbox = draw.textbbox((0, 0), line, font=font_main)
            text_width = bbox[2] - bbox[0]
            draw.text((width // 2 - text_width // 2, y), line, font=font_main, fill='#2F4F4F')
            y += 65

    # Место с выделением - УВЕЛИЧЕННЫЙ РАЗМЕР
    y += 100
    font_place = get_font(80, bold=True)  # Увеличен с 64 до 80

    if place == 1:
        place_text = "занявший I МЕСТО в олимпиаде"
        place_color = '#8B0000'
    elif place == 2:
        place_text = "занявшая II МЕСТО в олимпиаде"
        place_color = '#4169E1'
    elif place == 3:
        place_text = "занявший III МЕСТО в олимпиаде"
        place_color = '#CD853F'
    else:
        place_text = f"занявший {place} МЕСТО в олимпиаде"
        place_color = '#2F4F4F'

    bbox = draw.textbbox((0, 0), place_text, font=font_place)
    text_width = bbox[2] - bbox[0]
    draw.text((width // 2 - text_width // 2, y), place_text, font=font_place, fill=place_color)

    # Название олимпиады - БОЛЬШИМИ БУКВАМИ И УВЕЛИЧЕННЫЙ РАЗМЕР
    y += 140
    olympiad_line = f'«{olympiad_title.upper()}»'  # БОЛЬШИМИ БУКВАМИ
    font_olympiad = get_font(75, bold=True)  # Увеличен размер
    bbox = draw.textbbox((0, 0), olympiad_line, font=font_olympiad)
    text_width = bbox[2] - bbox[0]
    draw.text((width // 2 - text_width // 2, y), olympiad_line, font=font_olympiad, fill='#4169E1')

    # Результат
    if score is not None:
        y += 80

        # Заголовок результатов
        results_title = "и показал следующий результат:"
        bbox = draw.textbbox((0, 0), results_title, font=font_main)
        text_width = bbox[2] - bbox[0]
        draw.text((width // 2 - text_width // 2, y), results_title, font=font_main, fill='#000000')
        y += 100

        # Рамка для результатов
        box_width = 800
        box_height = 120
        box_x = width // 2 - box_width // 2
        box_y = y

        # Рисуем рамку
        draw.rectangle([box_x, box_y, box_x + box_width, box_y + box_height],
                       outline='#820000', width=4, fill='#FFF8DC')

        # Основной результат
        y += 30
        main_score_text = f"{score:.1f} баллов"
        font_score = get_font(60, bold=True)
        bbox = draw.textbbox((0, 0), main_score_text, font=font_score)
        text_width = bbox[2] - bbox[0]
        draw.text((width // 2 - text_width // 2, y), main_score_text, font=font_score, fill='#820000')

        y += 80

    # Дата в правый угол
    y += 200
    if hasattr(olympiad_end_date, 'strftime'):
        formatted_date = olympiad_end_date.strftime("«%d» %B %Y г.")
        # Заменяем английские названия месяцев на русские
        months = {
            'January': 'января', 'February': 'февраля', 'March': 'марта',
            'April': 'апреля', 'May': 'мая', 'June': 'июня',
            'July': 'июля', 'August': 'августа', 'September': 'сентября',
            'October': 'октября', 'November': 'ноября', 'December': 'декабря'
        }
        for eng, rus in months.items():
            formatted_date = formatted_date.replace(eng, rus)
    else:
        formatted_date = f"«___» _____________ {olympiad_end_date} г."

    date_font = get_font(36)
    bbox = draw.textbbox((0, 0), formatted_date, font=date_font)
    text_width = bbox[2] - bbox[0]
    date_x = width - text_width - 200
    date_y = y

    draw.text((date_x, date_y), formatted_date, font=date_font, fill='#2F4F4F')

    # Добавляем подписи
    img = add_signatures_to_certificate(img)

    return img

# Объединенный маршрут для скачивания сертификатов и дипломов

@app.route('/olympiad/<int:olympiad_id>/certificate')
@login_required
def download_certificate(olympiad_id):
    """Универсальный маршрут для скачивания сертификатов и дипломов"""
    olympiad = Olympiad.query.get_or_404(olympiad_id)
    cert_type = request.args.get('type', 'participation')

    # Проверяем участие пользователя
    participation = Participation.query.filter_by(
        user_id=current_user.id,
        olympiad_id=olympiad_id,
        status='completed'
    ).first()

    if not participation:
        flash('Вы не завершили эту олимпиаду', 'error')
        return redirect(url_for('view_olympiad', olympiad_id=olympiad_id))

    # Получаем информацию о специальности
    speciality_info = current_user.get_speciality_info()
    speciality = speciality_info['name'] if speciality_info else None

    # Для дипломов проверяем место и определяем параметры
    place = None
    score = participation.final_score  # Используем итоговый балл для всех типов
    time_bonus = participation.time_bonus if participation.time_bonus else 0
    filename_prefix = 'certificate_participation'

    if cert_type == 'winner':
        # Обновляем итоговые баллы
        update_all_final_scores(olympiad_id)

        # Определяем место пользователя
        rankings = Participation.query.filter_by(
            olympiad_id=olympiad_id,
            status='completed'
        ).order_by(Participation.final_score.desc()).all()

        user_place = None
        for i, p in enumerate(rankings, 1):
            if p.id == participation.id:
                user_place = i
                break

        # Проверяем, является ли пользователь призёром (топ-3)
        if user_place is None or user_place > 3:
            flash('Диплом победителя/призёра доступен только для участников, занявших 1-3 место', 'error')
            return redirect(url_for('olympiad_results', olympiad_id=olympiad_id))

        place = user_place
        place_names = {1: 'winner', 2: 'second', 3: 'third'}
        filename_prefix = f'diploma_{place_names.get(place, "prize")}'

    # Генерируем сертификат/диплом с баллами
    try:
        certificate_img = generate_certificate(
            user_name=current_user.full_name,
            olympiad_title=olympiad.title,
            olympiad_end_date=olympiad.end_time,
            certificate_type=cert_type,
            place=place,
            score=score,  # Передаем баллы для всех типов сертификатов
            time_bonus=time_bonus,  # Передаем временной бонус
            speciality=speciality
        )

        # Сохраняем в память
        img_io = BytesIO()
        certificate_img.save(img_io, 'PNG', quality=95, dpi=(300, 300))
        img_io.seek(0)

        # Формируем имя файла с указанием баллов
        score_text = f"{score:.1f}b"  # Добавляем баллы в имя файла
        filename = f'{filename_prefix}_{score_text}_{current_user.full_name}_{olympiad.title}_{datetime.now().strftime("%Y%m%d")}.png'
        filename = secure_filename(filename)

        return send_file(
            img_io,
            as_attachment=True,
            download_name=filename,
            mimetype='image/png'
        )

    except Exception as e:
        flash(f'Ошибка при создании {"диплома" if cert_type == "winner" else "сертификата"}: {str(e)}', 'error')
        return redirect(url_for('olympiad_results', olympiad_id=olympiad_id))
# Новые функции для расчета временного коэффициента
def calculate_time_bonus(actual_time, max_time, base_points):
    """
    Точный расчет временного бонуса на основе времени

    Формула: bonus = base_points * max_bonus_rate * (max_time - actual_time) / max_time

    Логика:
    - Мгновенное выполнение (0 сек) = максимальный бонус (25% от базовых баллов)
    - Выполнение в срок (100% времени) = 0 бонусов
    - Превышение времени = штраф (до -10% от базовых баллов)

    Args:
        actual_time: фактическое время выполнения в секундах
        max_time: максимальное время олимпиады в секундах
        base_points: базовые балл
        ы за правильные ответы

    Returns:
        float: точный временной бонус (может быть отрицательным)
    """

    if actual_time <= 0 or max_time <= 0 or base_points <= 0:
        return 0

    # Максимальный бонус за скорость (25% от базовых баллов)
    max_bonus_rate = 0.25

    # Максимальный штраф за превышение времени (10% от базовых баллов)
    max_penalty_rate = 0.10

    # Рассчитываем соотношение времени
    time_ratio = actual_time / max_time

    if time_ratio <= 1.0:
        # Выполнение в срок или быстрее - бонус
        # Формула: bonus = base_points * max_bonus_rate * (1 - time_ratio)
        time_bonus = base_points * max_bonus_rate * (1 - time_ratio)
    else:
        # Превышение времени - штраф
        # Штраф растет до максимума при превышении времени в 2 раза
        overtime_ratio = min(time_ratio - 1.0, 1.0)  # Ограничиваем максимальным штрафом
        time_bonus = -base_points * max_penalty_rate * overtime_ratio

    return round(time_bonus, 3)  # Точность до тысячных


def calculate_time_bonus_exponential(actual_time, max_time, base_points):
    """
    Экспоненциальная формула - более резкое падение бонуса

    Дает большой бонус за очень быстрое выполнение,
    но быстро убывает при увеличении времени
    """

    if actual_time <= 0 or max_time <= 0 or base_points <= 0:
        return 0

    max_bonus_rate = 0.30  # Увеличенный максимальный бонус
    time_ratio = actual_time / max_time

    if time_ratio <= 1.0:
        # Экспоненциальное убывание: bonus = base * max_bonus * exp(-3 * time_ratio)
        time_bonus = base_points * max_bonus_rate * math.exp(-3 * time_ratio)
    else:
        # Штраф за превышение времени
        overtime_ratio = min(time_ratio - 1.0, 1.0)
        time_bonus = -base_points * 0.15 * overtime_ratio

    return round(time_bonus, 3)


def calculate_time_bonus_logarithmic(actual_time, max_time, base_points):
    """
    Логарифмическая формула - плавное убывание бонуса

    Более справедливое распределение бонусов
    """

    if actual_time <= 0 or max_time <= 0 or base_points <= 0:
        return 0

    max_bonus_rate = 0.20
    time_ratio = actual_time / max_time

    if time_ratio <= 1.0:
        # Избегаем log(0), добавляем небольшое значение
        safe_ratio = max(time_ratio, 0.01)
        # Логарифмическая формула: bonus = base * max_bonus * -log(safe_ratio)
        time_bonus = base_points * max_bonus_rate * (-math.log(safe_ratio)) / 4.6  # нормализуем
    else:
        # Штраф за превышение времени
        overtime_ratio = min(time_ratio - 1.0, 1.0)
        time_bonus = -base_points * 0.12 * overtime_ratio

    return round(time_bonus, 3)


def get_time_performance_category_precise(actual_time, max_time):
    """
    Более точные категории производительности
    """
    if actual_time <= 0 or max_time <= 0:
        return "unknown", "⏰ Время не определено"

    time_percentage = (actual_time / max_time) * 100

    if time_percentage <= 10:
        return "lightning", "⚡ Молниеносно"
    elif time_percentage <= 25:
        return "excellent", "🚀 Превосходно"
    elif time_percentage <= 40:
        return "very_good", "⭐ Очень быстро"
    elif time_percentage <= 60:
        return "good", "✅ Быстро"
    elif time_percentage <= 80:
        return "normal", "⏱️ Нормально"
    elif time_percentage <= 100:
        return "slow", "🐌 Медленно"
    elif time_percentage <= 120:
        return "overtime", "⏰ Превышение времени"
    else:
        return "very_overtime", "🚨 Значительное превышение"


# Обновленная функция для расчета итогового балла
def calculate_final_score_precise(participation, early_finish=False):
    """
    Рассчитывает итоговый балл с точным учетом времени выполнения
    """
    if not participation.start_time or not participation.finish_time:
        participation.final_score = participation.total_points
        participation.duration_seconds = None
        participation.time_bonus = 0
        return

    # Получаем олимпиаду для расчета максимального времени
    olympiad = Olympiad.query.get(participation.olympiad_id)
    if not olympiad:
        participation.final_score = participation.total_points
        participation.time_bonus = 0
        return

    # Рассчитываем время выполнения в секундах
    duration = participation.finish_time - participation.start_time
    participation.duration_seconds = duration.total_seconds()

    # При досрочном завершении временной бонус не начисляется
    if early_finish:
        participation.time_bonus = 0
        participation.final_score = participation.total_points
        return

    # Максимальное время олимпиады в секундах
    max_duration = (olympiad.end_time - olympiad.start_time).total_seconds()

    # Выберите одну из формул (рекомендую линейную для начала):

    # 1. Линейная формула (рекомендуется)
    time_bonus = calculate_time_bonus(participation.duration_seconds, max_duration, participation.total_points)

    # 2. Экспоненциальная формула (для более агрессивного бонуса за скорость)
    # time_bonus = calculate_time_bonus_exponential(participation.duration_seconds, max_duration, participation.total_points)

    # 3. Логарифмическая формула (для плавного распределения)
    # time_bonus = calculate_time_bonus_logarithmic(participation.duration_seconds, max_duration, participation.total_points)

    # Сохраняем временной бонус отдельно для отображения
    participation.time_bonus = time_bonus

    # Итоговый балл = основные баллы + временной бонус
    participation.final_score = participation.total_points + time_bonus


# Демонстрация разных формул
def demo_time_bonus_formulas():
    """
    Демонстрирует работу разных формул расчета временного бонуса
    """
    print("=== ДЕМОНСТРАЦИЯ ФОРМУЛ ВРЕМЕННОГО БОНУСА ===")
    print("Базовые баллы: 100, Максимальное время: 3600 сек (1 час)")
    print()

    base_points = 100
    max_time = 3600  # 1 час

    test_times = [
        (180, "3 минуты (молниеносно)"),
        (900, "15 минут (очень быстро)"),
        (1800, "30 минут (быстро)"),
        (2700, "45 минут (нормально)"),
        (3600, "60 минут (в срок)"),
        (4500, "75 минут (превышение)"),
        (7200, "120 минут (большое превышение)")
    ]

    print(f"{'Время':<25} {'Линейная':<12} {'Экспонент.':<12} {'Логариф.':<12} {'Итого (лин.)':<12}")
    print("-" * 80)

    for actual_time, description in test_times:
        linear = calculate_time_bonus(actual_time, max_time, base_points)
        exponential = calculate_time_bonus_exponential(actual_time, max_time, base_points)
        logarithmic = calculate_time_bonus_logarithmic(actual_time, max_time, base_points)
        total_linear = base_points + linear

        print(f"{description:<25} {linear:>+8.2f} {exponential:>+8.2f} {logarithmic:>+8.2f} {total_linear:>8.2f}")


def get_time_performance_category(actual_time, max_time):
    """
    Определяет категорию производительности по времени
    """
    if actual_time <= 0 or max_time <= 0:
        return "unknown", "Время не определено"

    time_percentage = (actual_time / max_time) * 100

    if time_percentage <= 25:
        return "excellent", "⚡ Молниеносно"
    elif time_percentage <= 50:
        return "very_good", "🚀 Очень быстро"
    elif time_percentage <= 75:
        return "good", "⏱️ Быстро"
    elif time_percentage <= 100:
        return "normal", "✅ В срок"
    else:
        return "overtime", "⏰ Превышение времени"


# Добавляем функции в контекст шаблонов
@app.context_processor
def inject_time_functions():
    return dict(
        get_time_performance_category=get_time_performance_category,
        min=min,
        max=max
    )


def calculate_final_score(participation, early_finish=False):
    """
    Рассчитывает итоговый балл с учетом времени выполнения
    Новая формула: быстрое выполнение = больше бонусных баллов
    early_finish - если True, временной бонус не начисляется
    """
    if not participation.start_time or not participation.finish_time:
        participation.final_score = participation.total_points
        participation.duration_seconds = None
        participation.time_bonus = 0
        return

    # Получаем олимпиаду для расчета максимального времени
    olympiad = Olympiad.query.get(participation.olympiad_id)
    if not olympiad:
        participation.final_score = participation.total_points
        participation.time_bonus = 0
        return

    # Рассчитываем время выполнения в секундах
    duration = participation.finish_time - participation.start_time
    participation.duration_seconds = duration.total_seconds()

    # При досрочном завершении временной бонус не начисляется
    if early_finish:
        participation.time_bonus = 0
        participation.final_score = participation.total_points
        return

    # Максимальное время олимпиады в секундах
    max_duration = (olympiad.end_time - olympiad.start_time).total_seconds()

    # Расчет временного бонуса
    time_bonus = calculate_time_bonus(participation.duration_seconds, max_duration, participation.total_points)

    # Сохраняем временной бонус отдельно для отображения
    participation.time_bonus = time_bonus

    # Итоговый балл = основные баллы + временной бонус
    participation.final_score = participation.total_points + time_bonus


def update_all_final_scores(olympiad_id):
    """
    Обновляет итоговые баллы для всех завершенных участников олимпиады
    """
    participations = Participation.query.filter_by(
        olympiad_id=olympiad_id,
        status='completed'
    ).all()

    for participation in participations:
        calculate_final_score(participation)

    db.session.commit()


# Функция для пересчета баллов существующих участий
@app.route('/admin/fix_scoring/<int:olympiad_id>', methods=['GET'])
@login_required
def fix_scoring_system(olympiad_id):
    """Исправляет систему подсчета баллов для олимпиады"""
    if not current_user.is_admin:
        return jsonify({'success': False, 'message': 'Доступ запрещен'}), 403

    try:
        olympiad = Olympiad.query.get_or_404(olympiad_id)

        # Получаем все участия в этой олимпиаде
        participations = Participation.query.filter_by(olympiad_id=olympiad_id).all()

        fixed_count = 0
        for participation in participations:
            # Пересчитываем total_points на основе BlockResult
            block_results = BlockResult.query.filter_by(participation_id=participation.id).all()
            correct_total = sum(br.points_earned for br in block_results)

            if participation.total_points != correct_total:
                print(f"Участник {participation.user_id}: было {participation.total_points}, стало {correct_total}")
                participation.total_points = correct_total

                # Пересчитываем итоговый балл
                if participation.status == 'completed':
                    calculate_final_score(participation)

                fixed_count += 1

        db.session.commit()

        return jsonify({
            'success': True,
            'message': f'Исправлена система баллов для {fixed_count} участий в олимпиаде "{olympiad.title}"'
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Ошибка при исправлении: {str(e)}'
        }), 500


def recalculate_all_time_scores():
    """
    Пересчитывает временные коэффициенты для всех завершенных участий
    """
    completed_participations = Participation.query.filter_by(status='completed').all()

    for participation in completed_participations:
        if participation.start_time and participation.finish_time:
            calculate_final_score(participation)

    db.session.commit()
    return len(completed_participations)


# Models
class User(db.Model, UserMixin):
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(100), unique=True, nullable=False)
    password_hash = db.Column(db.String(200), nullable=False)
    full_name = db.Column(db.String(100), nullable=False)
    study_group = db.Column(db.String(50), nullable=True)
    course = db.Column(db.Integer, nullable=True)  # НОВОЕ ПОЛЕ
    speciality = db.Column(db.Text, nullable=True)  # JSON с информацией о специальности
    is_admin = db.Column(db.Boolean, default=False)
    participations = db.relationship('Participation', backref='user', lazy=True)

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

    def get_speciality_info(self):
        """Возвращает информацию о специальности пользователя"""
        if self.speciality:
            try:
                return json.loads(self.speciality)
            except:
                return None
        return None


class Olympiad(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(100), nullable=False)
    description = db.Column(db.Text, nullable=False)
    start_time = db.Column(db.DateTime, nullable=False)
    end_time = db.Column(db.DateTime, nullable=False)
    welcome_pdf = db.Column(db.String(200), nullable=True)
    created_by = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    blocks = db.relationship('Block', backref='olympiad', lazy=True, cascade="all, delete-orphan")
    participations = db.relationship('Participation', backref='olympiad', lazy=True, cascade="all, delete-orphan")


class Block(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    olympiad_id = db.Column(db.Integer, db.ForeignKey('olympiad.id'), nullable=False)
    title = db.Column(db.String(100), nullable=False)
    description = db.Column(db.Text, nullable=True)
    max_points = db.Column(db.Float, nullable=False)
    threshold_percentage = db.Column(db.Float, nullable=False)  # % для перехода на следующий блок
    order = db.Column(db.Integer, nullable=False)
    questions = db.relationship('Question', backref='block', lazy=True, cascade="all, delete-orphan")


class BlockResult(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    participation_id = db.Column(db.Integer, db.ForeignKey('participation.id'), nullable=False)
    block_id = db.Column(db.Integer, db.ForeignKey('block.id'), nullable=False)
    points_earned = db.Column(db.Float, default=0)
    completed_at = db.Column(db.DateTime, default=get_current_time)

    participation = db.relationship('Participation', backref='block_results')
    block = db.relationship('Block', backref='results')


class Question(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    block_id = db.Column(db.Integer, db.ForeignKey('block.id'), nullable=False)
    question_type = db.Column(db.String(20), nullable=False)  # 'test' или 'matching'
    text = db.Column(db.Text, nullable=False)
    options = db.Column(db.Text, nullable=True)  # JSON строка для вариантов ответа
    correct_answers = db.Column(db.Text, nullable=True)  # JSON строка для правильных ответов
    matches = db.Column(db.Text, nullable=True)  # JSON строка для пар соответствия
    points = db.Column(db.Float, nullable=False)
    answers = db.relationship('Answer', backref='question', lazy=True, cascade="all, delete-orphan")


class Participation(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    olympiad_id = db.Column(db.Integer, db.ForeignKey('olympiad.id'), nullable=False)
    start_time = db.Column(db.DateTime, nullable=True)
    finish_time = db.Column(db.DateTime, nullable=True)
    total_points = db.Column(db.Float, default=0)  # Основные баллы за правильные ответы
    final_score = db.Column(db.Float, default=0)  # Итоговый балл с учетом времени
    duration_seconds = db.Column(db.Float, nullable=True)  # Время выполнения в секундах
    time_bonus = db.Column(db.Float, default=0)  # Временной бонус отдельно
    status = db.Column(db.String(20), default='registered')  # 'registered', 'in_progress', 'completed'
    current_block = db.Column(db.Integer, nullable=True)
    answers = db.relationship('Answer', backref='participation', lazy=True, cascade="all, delete-orphan")


class Answer(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    participation_id = db.Column(db.Integer, db.ForeignKey('participation.id'), nullable=False)
    question_id = db.Column(db.Integer, db.ForeignKey('question.id'), nullable=False)
    answer_data = db.Column(db.Text, nullable=False)  # JSON строка для ответа
    is_correct = db.Column(db.Boolean, default=False)
    points_earned = db.Column(db.Float, default=0)
    answered_at = db.Column(db.DateTime, default=get_current_time)


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


@app.context_processor  # runs for every template render
def inject_now():
    # ИСПРАВЛЕНО: используем локальное время вместо UTC
    return {'now': get_current_time}


# Routes
@app.route('/')
def index():
    if current_user.is_authenticated:
        if current_user.is_admin:
            olympiads = Olympiad.query.all()
        else:
            # ИСПРАВЛЕНО: используем локальное время
            current_time = get_current_time()
            olympiads = Olympiad.query.filter(
                Olympiad.end_time > current_time
            ).all()
    else:
        olympiads = []
    return render_template('index.html', olympiads=olympiads)


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')

        user = User.query.filter_by(email=email).first()
        if user and user.check_password(password):
            login_user(user)
            next_page = request.args.get('next')
            return redirect(next_page or url_for('index'))
        flash('Неверный email или пароль', 'error')

    return render_template('login.html')


@app.route('/api/specialities', methods=['GET'])
def get_specialities():
    """API роут для получения списка специальностей"""
    try:
        response = requests.get('https://melsu.ru/api/specialities/list', timeout=10)
        if response.status_code == 200:
            return jsonify(response.json())
        else:
            return jsonify({'error': 'Не удалось получить список специальностей'}), 500
    except requests.RequestException:
        return jsonify({'error': 'Ошибка соединения с сервером'}), 500


@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        full_name = request.form.get('full_name')
        study_group = request.form.get('study_group')
        course = request.form.get('course')  # НОВОЕ ПОЛЕ
        speciality_id = request.form.get('speciality_id')

        if User.query.filter_by(email=email).first():
            flash('Email уже зарегистрирован', 'error')
            return redirect(url_for('register'))

        # Получаем информацию о специальности
        speciality_info = None
        if speciality_id:
            try:
                response = requests.get('https://melsu.ru/api/specialities/list', timeout=10)
                if response.status_code == 200:
                    specialities = response.json()
                    if speciality_id in specialities:
                        spec = specialities[speciality_id]
                        speciality_info = {
                            'id': spec['id'],
                            'spec_code': spec['spec_code'],
                            'name': spec['name'],
                            'department_name': spec['department_name'],
                            'faculty_name': spec['faculty_name'],
                            'faculty_acronym': spec['faculty_acronym'],
                            'level': spec['level']
                        }
            except requests.RequestException:
                flash('Не удалось сохранить информацию о специальности', 'warning')

        user = User(
            email=email,
            full_name=full_name,
            study_group=study_group,
            course=int(course) if course else None,  # НОВОЕ ПОЛЕ
            speciality=json.dumps(speciality_info) if speciality_info else None
        )
        user.set_password(password)

        db.session.add(user)
        db.session.commit()

        login_user(user)
        return redirect(url_for('index'))

    return render_template('register.html')


@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('index'))


@app.route('/profile')
@login_required
def profile():
    participations = Participation.query.filter_by(user_id=current_user.id).all()
    return render_template('profile.html', participations=participations)


# Admin routes
@app.route('/admin/olympiads', methods=['GET'])
@login_required
def admin_olympiads():
    if not current_user.is_admin:
        flash('У вас нет доступа к этой странице', 'error')
        return redirect(url_for('index'))

    olympiads = Olympiad.query.all()
    return render_template('admin/olympiads.html', olympiads=olympiads)


# Маршрут для управления пользователями
@app.route('/admin/users', methods=['GET'])
@login_required
def admin_users():
    if not current_user.is_admin:
        flash('У вас нет доступа к этой странице', 'error')
        return redirect(url_for('index'))

    users = User.query.all()
    return render_template('admin/users.html', users=users)


# Маршрут для аналитики
@app.route('/admin/analytics', methods=['GET'])
@login_required
def admin_analytics():
    if not current_user.is_admin:
        flash('У вас нет доступа к этой странице', 'error')
        return redirect(url_for('index'))

    # Общая статистика
    total_olympiads = Olympiad.query.count()
    total_users = User.query.count()
    total_participations = Participation.query.count()
    completed_participations = Participation.query.filter_by(status='completed').count()

    # Статистика по олимпиадам
    current_time = get_current_time()
    active_olympiads = Olympiad.query.filter(
        Olympiad.start_time <= current_time,
        Olympiad.end_time > current_time
    ).count()

    upcoming_olympiads = Olympiad.query.filter(
        Olympiad.start_time > current_time
    ).count()

    # Топ олимпиад по участникам
    olympiad_stats = db.session.query(
        Olympiad.title,
        db.func.count(Participation.id).label('participants')
    ).join(Participation).group_by(Olympiad.id).order_by(
        db.func.count(Participation.id).desc()
    ).limit(10).all()

    return render_template('admin/analytics.html',
                           total_olympiads=total_olympiads,
                           total_users=total_users,
                           total_participations=total_participations,
                           completed_participations=completed_participations,
                           active_olympiads=active_olympiads,
                           upcoming_olympiads=upcoming_olympiads,
                           olympiad_stats=olympiad_stats)


# Маршрут для настроек системы
@app.route('/admin/settings', methods=['GET'])
@login_required
def admin_settings():
    if not current_user.is_admin:
        flash('У вас нет доступа к этой странице', 'error')
        return redirect(url_for('index'))

    return render_template('admin/settings.html')


# Добавить роут для ручного пересчета временных коэффициентов
@app.route('/admin/recalculate_time_scores', methods=['POST'])
@login_required
def recalculate_time_scores():
    """Ручной пересчет временных коэффициентов"""
    if not current_user.is_admin:
        return jsonify({'success': False, 'message': 'Доступ запрещен'}), 403

    try:
        count = recalculate_all_time_scores()
        return jsonify({
            'success': True,
            'message': f'Пересчитаны временные коэффициенты для {count} участий'
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Ошибка при пересчете: {str(e)}'
        }), 500


# Маршрут для генерации DOCX документа с результатами
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


@app.route('/admin/olympiad/<int:olympiad_id>/export_docx', methods=['GET'])
@login_required
def export_rankings_docx(olympiad_id):
    # Проверка прав администратора
    if not current_user.is_admin:
        flash('У вас нет доступа к этой странице', 'error')
        return redirect(url_for('index'))

    olympiad = Olympiad.query.get_or_404(olympiad_id)
    update_all_final_scores(olympiad_id)
    participations = Participation.query.filter_by(
        olympiad_id=olympiad_id, status='completed'
    ).order_by(Participation.final_score.desc()).all()
    blocks = Block.query.filter_by(
        olympiad_id=olympiad_id
    ).order_by(Block.order).all()

    zip_io = BytesIO()
    with zipfile.ZipFile(zip_io, 'w', zipfile.ZIP_DEFLATED) as zipf:
        # Документы по этапам
        for block in blocks:
            doc = Document()
            font = doc.styles['Normal'].font
            font.name = 'Times New Roman';
            font.size = Pt(14)
            # Шапка
            for line in [
                'ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ',
                'ВЫСШЕГО ОБРАЗОВАНИЯ «МЕЛИТОПОЛЬСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ»',
                'Технический факультет',
                'кафедра «Гражданская безопасность»'
            ]:
                p = doc.add_paragraph();
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(line);
                r.font.name = 'Times New Roman';
                r.font.size = Pt(14);
                r.bold = True
            for _ in range(5): doc.add_paragraph()
            # Заголовок этапа
            p = doc.add_paragraph();
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f'ЭТАП {block.order}: {block.title}')
            r.font.name = 'Times New Roman';
            r.font.size = Pt(14);
            r.bold = True
            doc.add_paragraph()
            # Унифицированная таблица
            thr = block.max_points * (block.threshold_percentage / 100)
            rows = []
            for part in participations:
                br = BlockResult.query.filter_by(participation_id=part.id, block_id=block.id).first()
                if br and br.points_earned >= thr:
                    rows.append((part, br.points_earned))
            table = doc.add_table(rows=1, cols=5)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER;
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            cols = ['Место', 'ФИО студента', 'Группа', 'Направление подготовки', 'Баллы']
            for i, h in enumerate(cols):
                hdr[i].text = h;
                hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER;
                hdr[i].paragraphs[0].runs[0].bold = True
            for idx, (pt, pts) in enumerate(sorted(rows, key=lambda x: x[1], reverse=True), 1):
                usr = User.query.get(pt.user_id);
                spec = usr.get_speciality_info();
                spec = spec['name'] if spec else '-'
                row = table.add_row().cells
                vals = [str(idx), usr.full_name, usr.study_group or '-', spec, f"{pts:.1f}"]
                for j, c in enumerate(row): c.text = vals[j]; c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Подписи
            for _ in range(4): doc.add_paragraph()
            p = doc.add_paragraph();
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT;
            p.add_run(f'«___»____________ {datetime.now().year} г.').font.name = 'Times New Roman';
            p.runs[0].font.size = Pt(14)
            p = doc.add_paragraph();
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT;
            p.add_run('Члены жюри:').font.name = 'Times New Roman';
            p.runs[0].font.size = Pt(14);
            p.runs[0].bold = True
            doc.add_paragraph()
            jt = doc.add_table(rows=3, cols=4);
            jt.style = None
            widths = [Inches(1), Inches(1.5), Inches(1), Inches(3.5)]
            for ci, w in enumerate(widths):
                for rw in jt.rows: rw.cells[ci].width = w
            sigs = [['', '(подпись)', '', '(иниц., фам., степень, должность)']] * 3
            for ri, data in enumerate(sigs):
                rw = jt.rows[ri]
                for ci, txt in enumerate(data):
                    cell = rw.cells[ci]
                    if txt: cell.text = txt; cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            bio = BytesIO();
            doc.save(bio);
            bio.seek(0)
            zipf.writestr(f'etap_{block.order}.docx', bio.getvalue())
            # Документ ТОП-3
        doc = Document()
        font = doc.styles['Normal'].font
        font.name = 'Times New Roman';
        font.size = Pt(14)
        for line in [
            'ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ',
            'ВЫСШЕГО ОБРАЗОВАНИЯ «МЕЛИТОПОЛЬСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ»',
            'Технический факультет',
            'кафедра «Гражданская безопасность»'
        ]:
            p = doc.add_paragraph();
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line);
            r.font.name = 'Times New Roman';
            r.font.size = Pt(14);
            r.bold = True
        for _ in range(5): doc.add_paragraph()
        # Унифицированная таблица для ТОП-3
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER;
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        cols = ['Место', 'ФИО студента', 'Группа', 'Направление подготовки', 'Итоговый балл']
        for i, h in enumerate(cols): hdr[i].text = h; hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER;
        hdr[i].paragraphs[0].runs[0].bold = True
        for idx, part in enumerate(participations[:3], 1):
            usr = User.query.get(part.user_id)
            spec = usr.get_speciality_info();
            spec = spec['name'] if spec else '-'
            row = table.add_row().cells
            vals = [str(idx), usr.full_name, usr.study_group or '-', spec, f"{part.final_score:.1f}"]
            for j, c in enumerate(row): c.text = vals[j]; c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Подписи
        for _ in range(4): doc.add_paragraph()
        p = doc.add_paragraph();
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(f'«___»____________ {datetime.now().year} г.');
        r.font.name = 'Times New Roman';
        r.font.size = Pt(14)
        p = doc.add_paragraph();
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run('Члены жюри:');
        r.font.name = 'Times New Roman';
        r.font.size = Pt(14);
        r.bold = True
        doc.add_paragraph()
        jt = doc.add_table(rows=3, cols=4);
        jt.style = None
        widths = [Inches(1), Inches(1.5), Inches(1), Inches(3.5)]
        for ci, w in enumerate(widths):
            for rw in jt.rows: rw.cells[ci].width = w
        sigs = [['', '(подпись)', '', '(иниц., фам., степень, должность)']] * 3
        for ri, rowdata in enumerate(sigs):
            rw = jt.rows[ri]
            for ci, txt in enumerate(rowdata):
                cell = rw.cells[ci]
                if txt: cell.text = txt; cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        bio = BytesIO();
        doc.save(bio);
        bio.seek(0)
        zipf.writestr('top3.docx', bio.getvalue())
        # Документ всех участников
        doc = Document()
        font = doc.styles['Normal'].font
        font.name = 'Times New Roman';
        font.size = Pt(14)
        for line in [
            'ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ',
            'ВЫСШЕГО ОБРАЗОВАНИЯ «МЕЛИТОПОЛЬСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ»',
            'Технический факультет',
            'кафедра «Гражданская безопасность»'
        ]:
            p = doc.add_paragraph();
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line);
            r.font.name = 'Times New Roman';
            r.font.size = Pt(14);
            r.bold = True
        for _ in range(5): doc.add_paragraph()
        p = doc.add_paragraph();
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('СПИСОК ВСЕХ УЧАСТНИКОВ');
        r.font.name = 'Times New Roman';
        r.font.size = Pt(14);
        r.bold = True
        doc.add_paragraph()
        tbl = doc.add_table(rows=1, cols=5);
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER;
        tbl.style = 'Table Grid'
        hdr = tbl.rows[0].cells;
        cols = ['Место', 'ФИО', 'Группа', 'Направление подготовки', 'Итоговый балл']
        for i, c in enumerate(cols): hdr[i].text = c; hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER;
        hdr[i].paragraphs[0].runs[0].bold = True
        for idx, part in enumerate(participations, 1):
            usr = User.query.get(part.user_id);
            spec = usr.get_speciality_info();
            spec = spec['name'] if spec else '-'
            row = tbl.add_row().cells;
            vals = [str(idx), usr.full_name, usr.study_group or '-', spec, f"{part.final_score:.1f}"]
            for j, c in enumerate(row): c.text = vals[j]; c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Подписи
        for _ in range(4): doc.add_paragraph()
        p = doc.add_paragraph();
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(f'«___»____________ {datetime.now().year} г.')
        r.font.name = 'Times New Roman';
        r.font.size = Pt(14)
        p = doc.add_paragraph();
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run('Члены жюри:')
        r.font.name = 'Times New Roman';
        r.font.size = Pt(14);
        r.bold = True
        doc.add_paragraph()
        jt = doc.add_table(rows=3, cols=4);
        jt.style = None
        widths = [Inches(1), Inches(1.5), Inches(1), Inches(3.5)]
        for ci, w in enumerate(widths):
            for rw in jt.rows: rw.cells[ci].width = w
        sigs = [['', '(подпись)', '', '(иниц., фам., степень, должность)']] * 3
        for ri, rowdata in enumerate(sigs):
            rw = jt.rows[ri]
            for ci, txt in enumerate(rowdata):
                cell = rw.cells[ci]
                if txt: cell.text = txt; cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        bio = BytesIO();
        doc.save(bio);
        bio.seek(0)
        zipf.writestr('all_participants.docx', bio.getvalue())
    zip_io.seek(0)
    return send_file(zip_io, as_attachment=True,
                     download_name=f'stages_{olympiad.title}_{datetime.now().strftime("%Y%m%d_%H%M")}.zip',
                     mimetype='application/zip')


from flask import request, jsonify
import json
import re


@app.route('/admin/block/<int:block_id>/upload_questions', methods=['POST'])
@login_required
def upload_questions(block_id):
    """
    Загрузка вопросов для блока из файла
    Поддерживаемые форматы:
    1. Тесты:
       "1. Название вопроса" затем варианты ответа, правильные ответы начинаются с 4-х пробелов
    2. Сопоставление:
       "1. Название вопроса" затем пары для сопоставления в формате "Вариант 1 | Ответ 1"
    """
    if not current_user.is_admin:
        return jsonify({'success': False, 'message': 'Доступ запрещен'}), 403

    block = Block.query.get_or_404(block_id)

    # Проверяем, что файл есть в запросе
    if 'questions_file' not in request.files:
        return jsonify({'success': False, 'message': 'Файл не найден в запросе'})

    file = request.files['questions_file']
    if file.filename == '':
        return jsonify({'success': False, 'message': 'Файл не выбран'})

    # Читаем содержимое файла
    try:
        content = file.read().decode('utf-8')
    except UnicodeDecodeError:
        try:
            # Пробуем другую кодировку, если UTF-8 не работает
            file.seek(0)
            content = file.read().decode('windows-1251')
        except:
            return jsonify({'success': False,
                            'message': 'Не удалось прочитать файл. Проверьте кодировку (поддерживаются UTF-8 и Windows-1251)'})

    # Определяем тип блока по содержимому
    block_type = request.form.get('block_type')
    if not block_type:
        # Автоопределение типа блока по содержимому
        if '|' in content:
            block_type = 'matching'
        else:
            block_type = 'test'

    # Обработка содержимого в зависимости от типа блока
    questions_created = 0
    try:
        if block_type == 'test':
            questions_created = parse_test_questions(content, block_id)
        elif block_type == 'matching':
            questions_created = parse_matching_questions(content, block_id)
        else:
            return jsonify({'success': False, 'message': f'Неизвестный тип блока: {block_type}'})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Ошибка при обработке файла: {str(e)}'})

    # Обновляем равномерно баллы за все вопросы в блоке
    update_question_points(block_id)

    return jsonify({
        'success': True,
        'message': f'Успешно загружено {questions_created} вопросов в блок',
        'questions_count': questions_created
    })


def parse_test_questions(content, block_id):
    """Разбор содержимого файла с тестовыми вопросами"""
    lines = content.splitlines()

    questions = []
    current_question = None
    current_options = []
    current_correct = []

    for line in lines:
        line = line.rstrip()
        if not line:  # Пропускаем пустые строки
            continue

        # Новый вопрос начинается с номера и точки
        if re.match(r'^\d+\.', line):
            # Сохраняем предыдущий вопрос, если он есть
            if current_question:
                questions.append({
                    'text': current_question,
                    'options': current_options,
                    'correct_answers': current_correct
                })

            # Начинаем новый вопрос
            current_question = line.split('.', 1)[1].strip()
            current_options = []
            current_correct = []
        elif line.startswith('    '):  # Правильный ответ (4 пробела в начале)
            option = line.strip()
            if option not in current_options:
                current_options.append(option)
            current_correct.append(option)
        else:  # Обычный вариант ответа
            option = line.strip()
            if option and option not in current_options:
                current_options.append(option)

    # Добавляем последний вопрос
    if current_question:
        questions.append({
            'text': current_question,
            'options': current_options,
            'correct_answers': current_correct
        })

    # Сохраняем вопросы в базу данных
    questions_created = 0
    for q_data in questions:
        if not q_data['options'] or not q_data['correct_answers']:
            continue  # Пропускаем некорректные вопросы

        question = Question(
            block_id=block_id,
            question_type='test',
            text=q_data['text'],
            options=json.dumps(q_data['options']),
            correct_answers=json.dumps(q_data['correct_answers']),
            points=1.0  # Временное значение, будет обновлено позже
        )
        db.session.add(question)
        questions_created += 1

    db.session.commit()
    return questions_created


def parse_matching_questions(content, block_id):
    """Разбор содержимого файла с вопросами на сопоставление (обновленная версия для 3 колонок)"""
    lines = content.splitlines()

    questions = []
    current_question = None
    current_left_items = []
    current_middle_items = []
    current_right_items = []
    current_correct_matches = {}

    for line in lines:
        line = line.rstrip()
        if not line:  # Пропускаем пустые строки
            continue

        # Новый вопрос начинается с номера и точки
        if re.match(r'^\d+\.', line):
            # Сохраняем предыдущий вопрос, если он есть
            if current_question:
                questions.append({
                    'text': current_question,
                    'left_items': current_left_items,
                    'middle_items': current_middle_items,
                    'right_items': current_right_items,
                    'correct_matches': current_correct_matches
                })

            # Начинаем новый вопрос
            current_question = line.split('.', 1)[1].strip()
            current_left_items = []
            current_middle_items = []
            current_right_items = []
            current_correct_matches = {}
        elif '|' in line:  # Строка с парой для сопоставления
            parts = line.split('|')
            if len(parts) == 2:
                # Двухколоночное сопоставление (старый формат)
                left = parts[0].strip()
                right = parts[1].strip()
                if left and right:
                    if left not in current_left_items:
                        current_left_items.append(left)
                    if right not in current_right_items:
                        current_right_items.append(right)
                    current_correct_matches[left] = {'right': right}
            elif len(parts) == 3:
                # Трехколоночное сопоставление (новый формат)
                left = parts[0].strip()
                middle = parts[1].strip()
                right = parts[2].strip()
                if left and middle and right:
                    if left not in current_left_items:
                        current_left_items.append(left)
                    if middle not in current_middle_items:
                        current_middle_items.append(middle)
                    if right not in current_right_items:
                        current_right_items.append(right)
                    current_correct_matches[left] = {'middle': middle, 'right': right}
        elif line.startswith('M:'):  # Дополнительные средние элементы (отвлекающие)
            middle_item = line[2:].strip()
            if middle_item and middle_item not in current_middle_items:
                current_middle_items.append(middle_item)
        elif line.startswith('R:'):  # Дополнительные правые элементы (отвлекающие)
            right_item = line[2:].strip()
            if right_item and right_item not in current_right_items:
                current_right_items.append(right_item)

    # Добавляем последний вопрос
    if current_question:
        questions.append({
            'text': current_question,
            'left_items': current_left_items,
            'middle_items': current_middle_items,
            'right_items': current_right_items,
            'correct_matches': current_correct_matches
        })

    # Сохраняем вопросы в базу данных
    questions_created = 0
    for q_data in questions:
        if not q_data['left_items']:
            continue  # Пропускаем некорректные вопросы

        # Определяем тип сопоставления
        has_middle = len(q_data['middle_items']) > 0

        # Создаем новую структуру данных
        matches_data = {
            'left_items': q_data['left_items'],
            'middle_items': q_data['middle_items'] if has_middle else [],
            'right_items': q_data['right_items'],
            'correct_matches': q_data['correct_matches'],
            'columns': 3 if has_middle else 2  # Указываем количество колонок
        }

        question = Question(
            block_id=block_id,
            question_type='matching',
            text=q_data['text'],
            matches=json.dumps(matches_data),
            points=1.0  # Временное значение, будет обновлено позже
        )
        db.session.add(question)
        questions_created += 1

    db.session.commit()
    return questions_created


def update_question_points(block_id):
    """Обновление баллов за вопросы, чтобы их сумма равнялась max_points блока"""
    block = Block.query.get(block_id)
    questions = Question.query.filter_by(block_id=block_id).all()

    if not questions:
        return

    # Распределяем баллы поровну между всеми вопросами
    points_per_question = block.max_points / len(questions)

    for question in questions:
        question.points = points_per_question

    db.session.commit()



QUESTION_FILE_FORMAT = """
Формат файла для тестовых вопросов:
1. Название вопроса
Вариант ответа 1
Вариант ответа 2
    Правильный вариант ответа (начинается с 4 пробелов)
Вариант ответа 4

2. Еще один вопрос
Вариант ответа 1
    Правильный вариант 2
Вариант ответа 3

Формат файла для вопросов на сопоставление:
1. Название вопроса на сопоставление
Левая часть 1 | Правая часть 1
Левая часть 2 | Правая часть 2
Левая часть 3 | Правая часть 3

2. Еще один вопрос на сопоставление
Понятие 1 | Определение 1
Понятие 2 | Определение 2
"""


@app.route('/admin/block/<int:block_id>/file_format', methods=['GET'])
@login_required
def get_question_file_format(block_id):
    """Возвращает образец формата файла для загрузки вопросов"""
    if not current_user.is_admin:
        return jsonify({'success': False, 'message': 'Доступ запрещен'}), 403

    return jsonify({
        'success': True,
        'format': QUESTION_FILE_FORMAT
    })


def _get_month_name(month_num):
    """Возвращает название месяца на русском языке"""
    months = {
        1: 'января', 2: 'февраля', 3: 'марта', 4: 'апреля',
        5: 'мая', 6: 'июня', 7: 'июля', 8: 'августа',
        9: 'сентября', 10: 'октября', 11: 'ноября', 12: 'декабря'
    }
    return months.get(month_num, '')


@app.route('/admin/block/<int:block_id>/get_question', methods=['GET'])
@login_required
def get_question(block_id):
    if not current_user.is_admin:
        return jsonify({'success': False, 'message': 'У вас нет доступа к этой функции'}), 403

    question_id = request.args.get('question_id')
    if not question_id:
        return jsonify({'success': False, 'message': 'Не указан ID вопроса'}), 400

    question = Question.query.get_or_404(int(question_id))

    # Проверяем, принадлежит ли вопрос указанному блоку
    if question.block_id != block_id:
        return jsonify({'success': False, 'message': 'Вопрос не принадлежит указанному блоку'}), 403

    # Подготавливаем данные вопроса для отправки
    question_data = {
        'id': question.id,
        'text': question.text,
        'question_type': question.question_type,
        'points': question.points
    }

    # Добавляем специфичные для типа вопроса данные
    if question.question_type == 'test':
        question_data['options'] = question.options
        question_data['correct_answers'] = question.correct_answers

        # Для удобства работы с данными в JavaScript
        try:
            question_data['options_list'] = json.loads(question.options) if question.options else []
            question_data['correct_answers_list'] = json.loads(
                question.correct_answers) if question.correct_answers else []
        except json.JSONDecodeError:
            question_data['options_list'] = []
            question_data['correct_answers_list'] = []

    elif question.question_type == 'matching':
        question_data['matches'] = question.matches

        # Для удобства работы с данными в JavaScript
        try:
            question_data['matches_list'] = json.loads(question.matches) if question.matches else []
        except json.JSONDecodeError:
            question_data['matches_list'] = []

    return jsonify({
        'success': True,
        'question': question_data
    })


@app.route('/admin/block/<int:block_id>/update_question', methods=['POST'])
@login_required
def update_question(block_id):
    if not current_user.is_admin:
        return jsonify({'success': False, 'message': 'У вас нет доступа к этой функции'}), 403

    question_id = request.form.get('question_id')
    if not question_id:
        return jsonify({'success': False, 'message': 'Не указан ID вопроса'}), 400

    question = Question.query.get_or_404(int(question_id))

    # Проверяем, принадлежит ли вопрос указанному блоку
    if question.block_id != block_id:
        return jsonify({'success': False, 'message': 'Вопрос не принадлежит указанному блоку'}), 403

    # Обновляем общие поля
    question.text = request.form.get('text', question.text)

    # Обновляем специфичные для типа вопроса поля
    if question.question_type == 'test':
        options = request.form.getlist('options[]')
        correct_answers = request.form.getlist('correct_answers[]')

        if not options:
            return jsonify({'success': False, 'message': 'Необходимо указать хотя бы два варианта ответа'}), 400

        if len(options) < 2:
            return jsonify({'success': False, 'message': 'Необходимо указать хотя бы два варианта ответа'}), 400

        if not correct_answers:
            return jsonify({'success': False, 'message': 'Необходимо указать хотя бы один правильный ответ'}), 400

        # Убеждаемся, что все правильные ответы присутствуют в списке вариантов
        for answer in correct_answers:
            if answer not in options:
                return jsonify({'success': False, 'message': 'Правильный ответ должен быть в списке вариантов'}), 400

        question.options = json.dumps(options)
        question.correct_answers = json.dumps(correct_answers)

    elif question.question_type == 'matching':
        left_items = request.form.getlist('left_items[]')
        right_items = request.form.getlist('right_items[]')

        if not left_items or not right_items:
            return jsonify({'success': False, 'message': 'Необходимо указать хотя бы две пары для сопоставления'}), 400

        if len(left_items) != len(right_items):
            return jsonify(
                {'success': False, 'message': 'Количество элементов в левой и правой колонках должно совпадать'}), 400

        if len(left_items) < 2:
            return jsonify({'success': False, 'message': 'Необходимо указать хотя бы две пары для сопоставления'}), 400

        # Формируем пары
        matches = []
        for i in range(len(left_items)):
            matches.append({
                'left': left_items[i],
                'right': right_items[i]
            })

        question.matches = json.dumps(matches)

    # Сохраняем изменения
    try:
        db.session.commit()
        return jsonify({'success': True, 'message': 'Вопрос успешно обновлен'})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Ошибка при обновлении вопроса: {str(e)}")
        return jsonify({'success': False, 'message': f'Ошибка при обновлении вопроса: {str(e)}'}), 500


@app.route('/admin/block/<int:block_id>/delete_question', methods=['POST'])
@login_required
def delete_question(block_id):
    if not current_user.is_admin:
        return jsonify({'success': False, 'message': 'У вас нет доступа к этой функции'}), 403

    data = request.get_json()
    if not data or 'question_id' not in data:
        return jsonify({'success': False, 'message': 'Не указан ID вопроса'}), 400

    question_id = data['question_id']
    question = Question.query.get_or_404(int(question_id))

    # Проверяем, принадлежит ли вопрос указанному блоку
    if question.block_id != block_id:
        return jsonify({'success': False, 'message': 'Вопрос не принадлежит указанному блоку'}), 403

    # Удаляем вопрос
    try:
        db.session.delete(question)
        db.session.commit()

        # Пересчитываем баллы для оставшихся вопросов в блоке
        recalculate_points_for_block(block_id)

        return jsonify({'success': True, 'message': 'Вопрос успешно удален'})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Ошибка при удалении вопроса: {str(e)}")
        return jsonify({'success': False, 'message': f'Ошибка при удалении вопроса: {str(e)}'}), 500


def recalculate_points_for_block(block_id):
    """
    Пересчитывает баллы для всех вопросов в блоке,
    равномерно распределяя максимальное количество баллов блока.
    """
    block = Block.query.get_or_404(block_id)
    questions = Question.query.filter_by(block_id=block_id).all()

    if not questions:
        return

    # Равномерно распределяем баллы между всеми вопросами
    points_per_question = block.max_points / len(questions)

    for question in questions:
        question.points = points_per_question

    db.session.commit()

    # Пересчитываем баллы для всех существующих ответов
    for question in questions:
        answers = Answer.query.filter_by(question_id=question.id).all()
        for answer in answers:
            # Пересчитываем баллы для ответа
            if answer.is_correct:
                answer.points_earned = question.points
            else:
                # Для частично правильных ответов (например, matching)
                if question.question_type == 'matching' and answer.points_earned > 0:
                    # Сохраняем пропорцию правильности
                    old_proportion = answer.points_earned / question.points if question.points > 0 else 0
                    answer.points_earned = old_proportion * question.points

    db.session.commit()


@app.route('/admin/olympiad/<int:olympiad_id>/export_excel', methods=['GET'])
@login_required
def export_rankings_excel(olympiad_id):
    if not current_user.is_admin:
        flash('У вас нет доступа к этой странице', 'error')
        return redirect(url_for('index'))

    olympiad = Olympiad.query.get_or_404(olympiad_id)

    # Обновляем итоговые баллы перед экспортом
    update_all_final_scores(olympiad_id)

    # Получаем всех участников с результатами, сортируем по итоговому баллу
    participations = Participation.query.filter_by(
        olympiad_id=olympiad_id,
        status='completed'
    ).order_by(Participation.final_score.desc()).all()

    # Создаем workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Результаты"

    # Заголовки с временной информацией
    headers = ['Место', 'ФИО', 'Группа', 'Специальность', 'Баллы за задания',
               'Временной бонус', 'Итоговый балл', 'Время (мин)', 'Скорость', 'Начало', 'Завершение']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="820000", end_color="820000", fill_type="solid")
        cell.alignment = Alignment(horizontal="center")

    # Максимальное время олимпиады для расчета процентов
    olympiad_duration = (olympiad.end_time - olympiad.start_time).total_seconds()

    # Заполняем данными
    for row, participation in enumerate(participations, 2):
        user = User.query.get(participation.user_id)
        speciality_info = user.get_speciality_info()
        speciality_name = speciality_info['name'] if speciality_info else '-'

        duration = None
        speed_category = 'Неизвестно'
        if participation.duration_seconds:
            duration = participation.duration_seconds / 60
            time_percentage = (participation.duration_seconds / olympiad_duration) * 100

            # Определяем категорию скорости
            if time_percentage <= 25:
                speed_category = '⚡ Молниеносно'
            elif time_percentage <= 50:
                speed_category = '🚀 Очень быстро'
            elif time_percentage <= 75:
                speed_category = '⏱️ Быстро'
            elif time_percentage <= 100:
                speed_category = '✅ В срок'
            else:
                speed_category = '⏰ Превышение времени'

        time_bonus = participation.time_bonus if participation.time_bonus else 0

        data = [
            row - 1,  # Место
            user.full_name,
            user.study_group or '-',
            speciality_name,
            f"{participation.total_points:.2f}",  # Баллы за задания
            f"+{time_bonus:.2f}",  # Временной бонус
            f"{participation.final_score:.2f}",  # Итоговый балл
            f"{duration:.1f}" if duration else '-',
            speed_category,
            participation.start_time.strftime('%d.%m.%Y %H:%M') if participation.start_time else '-',
            participation.finish_time.strftime('%d.%m.%Y %H:%M') if participation.finish_time else '-'
        ]

        for col, value in enumerate(data, 1):
            cell = ws.cell(row=row, column=col, value=value)
            cell.alignment = Alignment(horizontal="center")

    # Автоподгонка ширины колонок
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width

    # Добавляем информацию об олимпиаде на отдельный лист
    info_ws = wb.create_sheet("Информация об олимпиаде")
    info_data = [
        ['Название олимпиады', olympiad.title],
        ['Описание', olympiad.description],
        ['Дата начала', olympiad.start_time.strftime('%d.%m.%Y %H:%M')],
        ['Дата окончания', olympiad.end_time.strftime('%d.%m.%Y %H:%M')],
        ['Всего участников', len(participations)],
        ['Дата экспорта', datetime.now().strftime('%d.%m.%Y %H:%M')],
        ['Применена система временных бонусов', 'Да'],
        ['', ''],
        ['Система временных бонусов:', ''],
        ['≤25% времени', '+20% от базовых баллов'],
        ['25-50% времени', '+10% от базовых баллов'],
        ['50-75% времени', '+5% от базовых баллов'],
        ['75-100% времени', '+1% от базовых баллов'],
        ['>100% времени', 'Нет бонуса'],
    ]

    for row, (key, value) in enumerate(info_data, 1):
        info_ws.cell(row=row, column=1, value=key).font = Font(bold=True)
        info_ws.cell(row=row, column=2, value=value)

    # Сохраняем в память
    output = BytesIO()
    wb.save(output)
    output.seek(0)

    filename = f'results_with_time_{olympiad.title}_{datetime.now().strftime("%Y%m%d_%H%M")}.xlsx'

    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )


@app.route('/admin/olympiad/<int:olympiad_id>/export_csv', methods=['GET'])
@login_required
def export_rankings_csv(olympiad_id):
    if not current_user.is_admin:
        flash('У вас нет доступа к этой странице', 'error')
        return redirect(url_for('index'))

    olympiad = Olympiad.query.get_or_404(olympiad_id)

    # Обновляем итоговые баллы перед экспортом
    update_all_final_scores(olympiad_id)

    # Получаем всех участников с результатами, сортируем по итоговому баллу
    participations = Participation.query.filter_by(
        olympiad_id=olympiad_id,
        status='completed'
    ).order_by(Participation.final_score.desc()).all()

    # Создаем CSV в памяти
    output = io.StringIO()
    writer = csv.writer(output, delimiter=';')

    # Заголовки с временной информацией
    writer.writerow(['Место', 'ФИО', 'Группа', 'Специальность', 'Баллы за задания',
                     'Временной бонус', 'Итоговый балл', 'Время (мин)', 'Скорость', 'Начало', 'Завершение'])

    # Максимальное время олимпиады
    olympiad_duration = (olympiad.end_time - olympiad.start_time).total_seconds()

    # Данные
    for i, participation in enumerate(participations, 1):
        user = User.query.get(participation.user_id)
        speciality_info = user.get_speciality_info()
        speciality_name = speciality_info['name'] if speciality_info else '-'

        duration = None
        speed_category = 'Неизвестно'
        if participation.duration_seconds:
            duration = participation.duration_seconds / 60
            time_percentage = (participation.duration_seconds / olympiad_duration) * 100

            if time_percentage <= 25:
                speed_category = 'Молниеносно'
            elif time_percentage <= 50:
                speed_category = 'Очень быстро'
            elif time_percentage <= 75:
                speed_category = 'Быстро'
            elif time_percentage <= 100:
                speed_category = 'В срок'
            else:
                speed_category = 'Превышение времени'

        time_bonus = participation.time_bonus if participation.time_bonus else 0

        writer.writerow([
            i,
            user.full_name,
            user.study_group or '-',
            speciality_name,
            f"{participation.total_points:.2f}",
            f"+{time_bonus:.2f}",
            f"{participation.final_score:.2f}",
            f"{duration:.1f}" if duration else '-',
            speed_category,
            participation.start_time.strftime('%d.%m.%Y %H:%M') if participation.start_time else '-',
            participation.finish_time.strftime('%d.%m.%Y %H:%M') if participation.finish_time else '-'
        ])

    output.seek(0)
    filename = f'results_with_time_{olympiad.title}_{datetime.now().strftime("%Y%m%d_%H%M")}.csv'

    return send_file(
        io.BytesIO(output.getvalue().encode('utf-8-sig')),
        as_attachment=True,
        download_name=filename,
        mimetype='text/csv'
    )


@app.route('/admin/olympiad/<int:olympiad_id>/export_detailed', methods=['GET'])
@login_required
def export_detailed_results(olympiad_id):
    """Детальный экспорт с результатами по блокам"""
    if not current_user.is_admin:
        flash('У вас нет доступа к этой странице', 'error')
        return redirect(url_for('index'))

    olympiad = Olympiad.query.get_or_404(olympiad_id)

    # Обновляем итоговые баллы перед экспортом
    update_all_final_scores(olympiad_id)

    # Получаем все блоки олимпиады
    blocks = Block.query.filter_by(olympiad_id=olympiad_id).order_by(Block.order).all()

    # Получаем всех участников, сортируем по итоговому баллу
    participations = Participation.query.filter_by(
        olympiad_id=olympiad_id,
        status='completed'
    ).order_by(Participation.final_score.desc()).all()

    # Создаем workbook с детальным анализом
    wb = Workbook()

    # Основной лист с результатами
    ws = wb.active
    ws.title = "Сводные результаты"

    # Формируем заголовки
    headers = ['Место', 'ФИО', 'Группа', 'Специальность', 'Баллы за задания', 'Временной бонус', 'Итоговый балл']
    for block in blocks:
        headers.append(f'Блок {block.order}: {block.title}')
    headers.extend(['Время (мин)', 'Скорость', 'Начало', 'Завершение'])

    # Записываем заголовки
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="820000", end_color="820000", fill_type="solid")
        cell.alignment = Alignment(horizontal="center")

    # Максимальное время олимпиады
    olympiad_duration = (olympiad.end_time - olympiad.start_time).total_seconds()

    # Заполняем данными
    for row, participation in enumerate(participations, 2):
        user = User.query.get(participation.user_id)
        speciality_info = user.get_speciality_info()
        speciality_name = speciality_info['name'] if speciality_info else '-'

        duration = None
        speed_category = 'Неизвестно'
        if participation.duration_seconds:
            duration = participation.duration_seconds / 60
            time_percentage = (participation.duration_seconds / olympiad_duration) * 100

            if time_percentage <= 25:
                speed_category = 'Молниеносно'
            elif time_percentage <= 50:
                speed_category = 'Очень быстро'
            elif time_percentage <= 75:
                speed_category = 'Быстро'
            elif time_percentage <= 100:
                speed_category = 'В срок'
            else:
                speed_category = 'Превышение времени'

        time_bonus = participation.time_bonus if participation.time_bonus else 0

        # Основные данные
        data = [
            row - 1,  # Место
            user.full_name,
            user.study_group or '-',
            speciality_name,
            f"{participation.total_points:.2f}",  # Баллы за задания
            f"+{time_bonus:.2f}",  # Временной бонус
            f"{participation.final_score:.2f}"  # Итоговый балл
        ]

        # Баллы по блокам
        for block in blocks:
            block_result = BlockResult.query.filter_by(
                participation_id=participation.id,
                block_id=block.id
            ).first()

            if block_result:
                data.append(f"{block_result.points_earned:.1f}")
            else:
                # Подсчитываем из ответов, если нет записи в BlockResult
                questions = Question.query.filter_by(block_id=block.id).all()
                answers = Answer.query.filter(
                    Answer.participation_id == participation.id,
                    Answer.question_id.in_([q.id for q in questions])
                ).all()

                if answers:
                    total_points = sum(answer.points_earned for answer in answers)
                    data.append(f"{total_points:.1f}")
                else:
                    data.append("0.0")

        # Время и статус
        data.extend([
            f"{duration:.1f}" if duration else '-',
            speed_category,
            participation.start_time.strftime('%d.%m.%Y %H:%M') if participation.start_time else '-',
            participation.finish_time.strftime('%d.%m.%Y %H:%M') if participation.finish_time else '-'
        ])

        for col, value in enumerate(data, 1):
            cell = ws.cell(row=row, column=col, value=value)
            cell.alignment = Alignment(horizontal="center")

    # Автоподгонка ширины
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width

    # Добавляем статистику по блокам
    stats_ws = wb.create_sheet("Статистика по блокам")

    # Заголовки статистики
    stats_headers = ['Блок', 'Средний балл', 'Максимум', 'Минимум', 'Прошли порог', 'Процент прохождения']
    for col, header in enumerate(stats_headers, 1):
        cell = stats_ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True)

    # Статистика по каждому блоку
    for row, block in enumerate(blocks, 2):
        block_results = BlockResult.query.filter_by(block_id=block.id).all()

        if block_results:
            points = [br.points_earned for br in block_results]
            avg_points = sum(points) / len(points)
            max_points = max(points)
            min_points = min(points)

            # Считаем сколько прошли порог
            threshold_points = block.max_points * (block.threshold_percentage / 100)
            passed_threshold = len([p for p in points if p >= threshold_points])
            pass_percentage = (passed_threshold / len(points)) * 100
        else:
            avg_points = max_points = min_points = 0
            passed_threshold = 0
            pass_percentage = 0

        stats_data = [
            f'Блок {block.order}: {block.title}',
            f"{avg_points:.1f}",
            f"{max_points:.1f}",
            f"{min_points:.1f}",
            f"{passed_threshold}/{len(block_results) if block_results else 0}",
            f"{pass_percentage:.1f}%"
        ]

        for col, value in enumerate(stats_data, 1):
            stats_ws.cell(row=row, column=col, value=value)

    # Сохраняем файл
    output = BytesIO()
    wb.save(output)
    output.seek(0)

    filename = f'detailed_results_with_time_{olympiad.title}_{datetime.now().strftime("%Y%m%d_%H%M")}.xlsx'

    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )


def get_month_name(month_num):
    """Возвращает название месяца на русском языке"""
    months = {
        1: 'января', 2: 'февраля', 3: 'марта', 4: 'апреля',
        5: 'мая', 6: 'июня', 7: 'июля', 8: 'августа',
        9: 'сентября', 10: 'октября', 11: 'ноября', 12: 'декабря'
    }
    return months.get(month_num, 'месяца')


# Маршрут для изменения статуса пользователя
@app.route('/admin/users/<int:user_id>/toggle_admin', methods=['POST'])
@login_required
def toggle_user_admin(user_id):
    if not current_user.is_admin:
        return jsonify({'success': False, 'message': 'Доступ запрещен'}), 403

    user = User.query.get_or_404(user_id)

    # Защита от отключения админки у самого себя
    if user.id == current_user.id:
        return jsonify({'success': False, 'message': 'Нельзя изменить собственный статус администратора'})

    user.is_admin = not user.is_admin
    db.session.commit()

    status = 'добавлены' if user.is_admin else 'отозваны'
    return jsonify({'success': True, 'message': f'Права администратора {status}'})


# Маршрут для удаления пользователя
@app.route('/admin/users/<int:user_id>/delete', methods=['POST'])
@login_required
def delete_user(user_id):
    if not current_user.is_admin:
        return jsonify({'success': False, 'message': 'Доступ запрещен'}), 403

    user = User.query.get_or_404(user_id)

    # Защита от удаления самого себя
    if user.id == current_user.id:
        return jsonify({'success': False, 'message': 'Нельзя удалить собственного пользователя'})

    # Удаляем связанные участия в олимпиадах
    Participation.query.filter_by(user_id=user.id).delete()

    db.session.delete(user)
    db.session.commit()

    return jsonify({'success': True, 'message': 'Пользователь успешно удален'})


@app.route('/admin/olympiad/create', methods=['POST'])
@login_required
def create_olympiad():
    if not current_user.is_admin:
        return jsonify({'success': False, 'message': 'Доступ запрещен'}), 403

    title = request.form.get('title')
    description = request.form.get('description')

    # ИСПРАВЛЕНО: парсим время как локальное
    start_time = datetime.strptime(request.form.get('start_time'), '%Y-%m-%dT%H:%M')
    end_time = datetime.strptime(request.form.get('end_time'), '%Y-%m-%dT%H:%M')

    pdf_file = request.files.get('welcome_pdf')
    welcome_pdf = None

    if pdf_file and pdf_file.filename:
        filename = secure_filename(f"{uuid.uuid4()}_{pdf_file.filename}")
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        pdf_file.save(pdf_path)
        welcome_pdf = filename

    olympiad = Olympiad(
        title=title,
        description=description,
        start_time=start_time,
        end_time=end_time,
        welcome_pdf=welcome_pdf,
        created_by=current_user.id
    )

    db.session.add(olympiad)
    db.session.commit()

    return jsonify({'success': True, 'olympiad_id': olympiad.id})


@app.route('/admin/olympiad/<int:olympiad_id>', methods=['GET'])
@login_required
def edit_olympiad(olympiad_id):
    if not current_user.is_admin:
        flash('У вас нет доступа к этой странице', 'error')
        return redirect(url_for('index'))

    olympiad = Olympiad.query.get_or_404(olympiad_id)
    blocks = Block.query.filter_by(olympiad_id=olympiad_id).order_by(Block.order).all()

    return render_template('admin/edit_olympiad.html', olympiad=olympiad, blocks=blocks)


@app.route('/admin/olympiad/<int:olympiad_id>/update', methods=['POST'])
@login_required
def update_olympiad(olympiad_id):
    if not current_user.is_admin:
        return jsonify({'success': False, 'message': 'Доступ запрещен'}), 403

    olympiad = Olympiad.query.get_or_404(olympiad_id)

    olympiad.title = request.form.get('title')
    olympiad.description = request.form.get('description')

    # ИСПРАВЛЕНО: парсим время как локальное
    olympiad.start_time = datetime.strptime(request.form.get('start_time'), '%Y-%m-%dT%H:%M')
    olympiad.end_time = datetime.strptime(request.form.get('end_time'), '%Y-%m-%dT%H:%M')

    pdf_file = request.files.get('welcome_pdf')
    if pdf_file and pdf_file.filename:
        # Удаляем старый файл, если он есть
        if olympiad.welcome_pdf:
            old_path = os.path.join(app.config['UPLOAD_FOLDER'], olympiad.welcome_pdf)
            if os.path.exists(old_path):
                os.remove(old_path)

        filename = secure_filename(f"{uuid.uuid4()}_{pdf_file.filename}")
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        pdf_file.save(pdf_path)
        olympiad.welcome_pdf = filename

    db.session.commit()

    return jsonify({'success': True})


@app.route('/admin/olympiad/<int:olympiad_id>/add_block', methods=['POST'])
@login_required
def add_block(olympiad_id):
    if not current_user.is_admin:
        return jsonify({'success': False, 'message': 'Доступ запрещен'}), 403

    olympiad = Olympiad.query.get_or_404(olympiad_id)

    # Получаем максимальный order для блоков в олимпиаде
    max_order = db.session.query(db.func.max(Block.order)).filter(Block.olympiad_id == olympiad_id).scalar() or 0

    title = request.form.get('title')
    description = request.form.get('description')
    max_points = float(request.form.get('max_points'))
    threshold_percentage = float(request.form.get('threshold_percentage'))

    block = Block(
        olympiad_id=olympiad_id,
        title=title,
        description=description,
        max_points=max_points,
        threshold_percentage=threshold_percentage,
        order=max_order + 1
    )

    db.session.add(block)
    db.session.commit()

    return jsonify({'success': True, 'block_id': block.id})


@app.route('/admin/block/<int:block_id>/edit', methods=['GET'])
@login_required
def edit_block(block_id):
    if not current_user.is_admin:
        flash('У вас нет доступа к этой странице', 'error')
        return redirect(url_for('index'))

    block = Block.query.get_or_404(block_id)
    questions = Question.query.filter_by(block_id=block_id).all()

    # Предварительная обработка данных вопросов
    questions = prepare_question_data(questions)

    return render_template('admin/edit_block.html', block=block, questions=questions)


# Исправленная функция для создания блока с правильным распределением баллов
@app.route('/admin/block/<int:block_id>/add_question', methods=['POST'])
@login_required
def add_question(block_id):
    if not current_user.is_admin:
        return jsonify({'success': False, 'message': 'Доступ запрещен'}), 403

    block = Block.query.get_or_404(block_id)

    question_type = request.form.get('question_type')
    text = request.form.get('text')

    if question_type == 'test':
        options = request.form.getlist('options[]')
        correct_answers = request.form.getlist('correct_answers[]')

        question = Question(
            block_id=block_id,
            question_type=question_type,
            text=text,
            options=json.dumps(options),
            correct_answers=json.dumps(correct_answers),
            points=1.0  # Временное значение
        )
    elif question_type == 'matching':
        left_items = request.form.getlist('left_items[]')
        middle_items = request.form.getlist('middle_items[]')
        right_items = request.form.getlist('right_items[]')

        # Определяем количество колонок
        has_middle = len(middle_items) > 0 and any(item.strip() for item in middle_items)
        columns = 3 if has_middle else 2

        correct_matches_data = {}

        # Получаем соответствия из формы
        for i, left_item in enumerate(left_items):
            if has_middle:
                # Трехколоночное сопоставление
                middle_key = f'match_middle_{i}'
                right_key = f'match_right_{i}'
                if middle_key in request.form and right_key in request.form:
                    correct_matches_data[left_item] = {
                        'middle': request.form[middle_key],
                        'right': request.form[right_key]
                    }
            else:
                # Двухколоночное сопоставление
                match_key = f'match_{i}'
                if match_key in request.form:
                    correct_matches_data[left_item] = {
                        'right': request.form[match_key]
                    }

        # Создаем структуру данных
        matches_data = {
            'left_items': left_items,
            'middle_items': middle_items if has_middle else [],
            'right_items': right_items,
            'correct_matches': correct_matches_data,
            'columns': columns
        }

        question = Question(
            block_id=block_id,
            question_type=question_type,
            text=text,
            matches=json.dumps(matches_data),
            points=1.0  # Временное значение
        )

    db.session.add(question)
    db.session.commit()

    # Пересчитываем баллы для всех вопросов в блоке
    update_question_points(block_id)

    return jsonify({'success': True, 'question_id': question.id})

# Дополнительная функция для проверки и исправления баллов
@app.route('/admin/fix_points/<int:olympiad_id>', methods=['POST'])
@login_required
def fix_olympiad_points(olympiad_id):
    """Исправляет баллы для всех блоков олимпиады"""
    if not current_user.is_admin:
        return jsonify({'success': False, 'message': 'Доступ запрещен'}), 403

    try:
        olympiad = Olympiad.query.get_or_404(olympiad_id)
        blocks = Block.query.filter_by(olympiad_id=olympiad_id).all()

        fixed_blocks = 0
        for block in blocks:
            questions = Question.query.filter_by(block_id=block.id).all()
            if questions:
                # Пересчитываем баллы
                recalculate_points_for_block(block.id)
                fixed_blocks += 1

        # Пересчитываем итоговые баллы для всех участников
        update_all_final_scores(olympiad_id)

        return jsonify({
            'success': True,
            'message': f'Исправлены баллы для {fixed_blocks} блоков олимпиады "{olympiad.title}"'
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Ошибка при исправлении баллов: {str(e)}'
        }), 500


# Маршрут для диагностики блока
@app.route('/admin/block/<int:block_id>/diagnose', methods=['GET'])
@login_required
def diagnose_block(block_id):
    """Диагностика блока для проверки правильности распределения баллов"""
    if not current_user.is_admin:
        return jsonify({'success': False, 'message': 'Доступ запрещен'}), 403

    try:
        block = Block.query.get_or_404(block_id)
        questions = Question.query.filter_by(block_id=block_id).all()

        # Подсчитываем статистику
        total_question_points = sum(q.points for q in questions)
        points_per_question = block.max_points / len(questions) if questions else 0

        # Проверяем участников с ответами
        participants_with_answers = db.session.query(Answer.participation_id).filter(
            Answer.question_id.in_([q.id for q in questions])
        ).distinct().count() if questions else 0

        return jsonify({
            'success': True,
            'block_name': block.title,
            'max_points': block.max_points,
            'questions_count': len(questions),
            'points_per_question': round(points_per_question, 2),
            'total_question_points': round(total_question_points, 2),
            'participants_count': participants_with_answers,
            'questions_details': [{
                'id': q.id,
                'text': q.text[:50] + '...' if len(q.text) > 50 else q.text,
                'points': q.points,
                'type': q.question_type
            } for q in questions[:5]]  # Показываем первые 5 вопросов
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Ошибка диагностики: {str(e)}'
        }), 500


# Маршрут для полной диагностики олимпиады
@app.route('/admin/olympiad/<int:olympiad_id>/diagnose', methods=['GET'])
@login_required
def diagnose_olympiad(olympiad_id):
    """Полная диагностика олимпиады"""
    if not current_user.is_admin:
        return jsonify({'success': False, 'message': 'Доступ запрещен'}), 403

    try:
        olympiad = Olympiad.query.get_or_404(olympiad_id)
        blocks = Block.query.filter_by(olympiad_id=olympiad_id).order_by(Block.order).all()

        blocks_info = []
        total_olympiad_points = 0

        for block in blocks:
            questions = Question.query.filter_by(block_id=block.id).all()
            total_question_points = sum(q.points for q in questions)

            # Проверяем участников
            participants_count = db.session.query(Answer.participation_id).filter(
                Answer.question_id.in_([q.id for q in questions])
            ).distinct().count() if questions else 0

            block_info = {
                'id': block.id,
                'title': block.title,
                'order': block.order,
                'max_points': block.max_points,
                'questions_count': len(questions),
                'total_question_points': round(total_question_points, 2),
                'participants_count': participants_count,
                'is_correct': abs(total_question_points - block.max_points) < 0.01
            }

            blocks_info.append(block_info)
            total_olympiad_points += block.max_points

        # Проверяем участников олимпиады
        total_participants = Participation.query.filter_by(olympiad_id=olympiad_id).count()
        completed_participants = Participation.query.filter_by(
            olympiad_id=olympiad_id,
            status='completed'
        ).count()

        return jsonify({
            'success': True,
            'olympiad_title': olympiad.title,
            'total_olympiad_points': total_olympiad_points,
            'blocks_count': len(blocks),
            'total_participants': total_participants,
            'completed_participants': completed_participants,
            'blocks': blocks_info,
            'has_issues': any(not block['is_correct'] for block in blocks_info)
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Ошибка диагностики: {str(e)}'
        }), 500


# Маршрут для быстрого исправления одного блока
@app.route('/admin/block/<int:block_id>/fix_points', methods=['POST'])
@login_required
def fix_block_points(block_id):
    """Быстрое исправление баллов для одного блока"""
    if not current_user.is_admin:
        return jsonify({'success': False, 'message': 'Доступ запрещен'}), 403

    try:
        block = Block.query.get_or_404(block_id)

        # Исправляем баллы
        recalculate_points_for_block(block_id)

        # Пересчитываем итоговые баллы участников
        olympiad_id = block.olympiad_id
        update_all_final_scores(olympiad_id)

        return jsonify({
            'success': True,
            'message': f'Баллы для блока "{block.title}" успешно исправлены'
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Ошибка при исправлении: {str(e)}'
        }), 500


# Дополнительный маршрут для показа детальной информации об участнике
@app.route('/admin/participation/<int:participation_id>/details', methods=['GET'])
@login_required
def participation_details(participation_id):
    """Детальная информация об участии"""
    if not current_user.is_admin:
        return jsonify({'success': False, 'message': 'Доступ запрещен'}), 403

    try:
        participation = Participation.query.get_or_404(participation_id)
        user = User.query.get(participation.user_id)
        olympiad = Olympiad.query.get(participation.olympiad_id)

        # Получаем результаты по блокам
        blocks_results = []
        blocks = Block.query.filter_by(olympiad_id=olympiad.id).order_by(Block.order).all()

        for block in blocks:
            block_result = BlockResult.query.filter_by(
                participation_id=participation.id,
                block_id=block.id
            ).first()

            # Альтернативный подсчет через ответы
            questions = Question.query.filter_by(block_id=block.id).all()
            answers = Answer.query.filter(
                Answer.participation_id == participation.id,
                Answer.question_id.in_([q.id for q in questions])
            ).all()

            points_from_answers = sum(answer.points_earned for answer in answers)

            blocks_results.append({
                'block_title': block.title,
                'block_order': block.order,
                'max_points': block.max_points,
                'result_points': block_result.points_earned if block_result else 0,
                'answers_points': points_from_answers,
                'questions_answered': len(answers),
                'total_questions': len(questions),
                'completed': block_result is not None
            })

        return jsonify({
            'success': True,
            'user_name': user.full_name,
            'user_email': user.email,
            'olympiad_title': olympiad.title,
            'status': participation.status,
            'total_points': participation.total_points,
            'final_score': participation.final_score,
            'time_bonus': participation.time_bonus,
            'start_time': participation.start_time.isoformat() if participation.start_time else None,
            'finish_time': participation.finish_time.isoformat() if participation.finish_time else None,
            'blocks_results': blocks_results
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Ошибка получения информации: {str(e)}'
        }), 500

@app.route('/olympiad/<int:olympiad_id>/view', methods=['GET'])
@login_required
def view_olympiad(olympiad_id):
    olympiad = Olympiad.query.get_or_404(olympiad_id)

    # ИСПРАВЛЕНО: используем локальное время для проверки доступности
    current_time = get_current_time()
    if not current_user.is_admin and olympiad.end_time < current_time:
        flash('Олимпиада завершена', 'error')
        return redirect(url_for('index'))

    # Проверяем, зарегистрирован ли пользователь на эту олимпиаду
    participation = Participation.query.filter_by(
        user_id=current_user.id,
        olympiad_id=olympiad_id
    ).first()

    return render_template('olympiad/view.html', olympiad=olympiad, participation=participation)


@app.route('/olympiad/<int:olympiad_id>/register', methods=['POST'])
@login_required
def register_olympiad(olympiad_id):
    olympiad = Olympiad.query.get_or_404(olympiad_id)

    # Проверяем, не зарегистрирован ли пользователь уже
    existing = Participation.query.filter_by(
        user_id=current_user.id,
        olympiad_id=olympiad_id
    ).first()

    if existing:
        return jsonify({'success': False, 'message': 'Вы уже зарегистрированы на эту олимпиаду'})

    participation = Participation(
        user_id=current_user.id,
        olympiad_id=olympiad_id,
        status='registered'
    )

    db.session.add(participation)
    db.session.commit()

    return jsonify({'success': True})


@app.route('/olympiad/<int:olympiad_id>/start', methods=['POST'])
@login_required
def start_olympiad(olympiad_id):
    olympiad = Olympiad.query.get_or_404(olympiad_id)

    # Проверяем, не начата ли уже олимпиада
    participation = Participation.query.filter_by(
        user_id=current_user.id,
        olympiad_id=olympiad_id
    ).first()

    if not participation:
        return jsonify({'success': False, 'message': 'Вы не зарегистрированы на эту олимпиаду'})

    if participation.status == 'in_progress':
        return jsonify({'success': True, 'redirect': url_for('take_olympiad', olympiad_id=olympiad_id)})

    if participation.status == 'completed':
        return jsonify({'success': False, 'message': 'Вы уже завершили эту олимпиаду'})

    # ИСПРАВЛЕНО: используем локальное время для проверки времени начала
    current_time = get_current_time()
    if current_time < olympiad.start_time:
        return jsonify({
            'success': False,
            'message': f'Олимпиада начнется {olympiad.start_time.strftime("%d.%m.%Y в %H:%M")}'
        })

    if current_time > olympiad.end_time:
        return jsonify({'success': False, 'message': 'Время олимпиады истекло'})

    # Ищем первый блок
    first_block = Block.query.filter_by(olympiad_id=olympiad_id, order=1).first()
    if not first_block:
        return jsonify({'success': False, 'message': 'Олимпиада не содержит блоков'})

    participation.status = 'in_progress'
    participation.start_time = current_time
    participation.current_block = first_block.id

    db.session.commit()

    return jsonify({'success': True, 'redirect': url_for('take_olympiad', olympiad_id=olympiad_id)})


@app.route('/olympiad/<int:olympiad_id>/take', methods=['GET'])
@login_required
def take_olympiad(olympiad_id):
    olympiad = Olympiad.query.get_or_404(olympiad_id)

    # Проверяем участие пользователя
    participation = Participation.query.filter_by(
        user_id=current_user.id,
        olympiad_id=olympiad_id,
        status='in_progress'
    ).first()

    if not participation:
        flash('Вы не участвуете в этой олимпиаде или она уже завершена', 'error')
        return redirect(url_for('view_olympiad', olympiad_id=olympiad_id))

    # Получаем текущий блок
    current_block = Block.query.get(participation.current_block)
    if not current_block:
        flash('Ошибка: блок не найден', 'error')
        return redirect(url_for('view_olympiad', olympiad_id=olympiad_id))

    # Получаем вопросы текущего блока
    questions = Question.query.filter_by(block_id=current_block.id).all()

    # Подготавливаем данные вопросов для корректного отображения
    questions = prepare_question_data(questions)

    # Получаем ответы пользователя на вопросы этого блока
    user_answers = {}
    for question in questions:
        answer = Answer.query.filter_by(
            participation_id=participation.id,
            question_id=question.id
        ).first()
        if answer:
            try:
                user_answers[question.id] = json.loads(answer.answer_data)
            except:
                user_answers[question.id] = []

    return render_template(
        'olympiad/take.html',
        olympiad=olympiad,
        block=current_block,
        questions=questions,
        user_answers=user_answers,
        participation=participation
    )


# 3. Update the submit_answer route to handle three-column matching
# Исправленная функция submit_answer - НЕ обновляем total_points здесь
@app.route('/olympiad/<int:olympiad_id>/submit_answer', methods=['POST'])
@login_required
def submit_answer(olympiad_id):
    data = request.get_json()
    question_id = data.get('question_id')
    answer_data = data.get('answer_data')

    question = Question.query.get_or_404(question_id)

    # Проверяем участие пользователя
    participation = Participation.query.filter_by(
        user_id=current_user.id,
        olympiad_id=olympiad_id,
        status='in_progress'
    ).first()

    if not participation:
        return jsonify({'success': False, 'message': 'Вы не участвуете в этой олимпиаде'}), 403

    # Проверяем, относится ли вопрос к текущему блоку
    if question.block_id != participation.current_block:
        return jsonify({'success': False, 'message': 'Вопрос не принадлежит текущему блоку'}), 403

    # Проверяем правильность ответа
    is_correct = False
    points_earned = 0

    if question.question_type == 'test':
        correct_answers = set(json.loads(question.correct_answers))
        user_answers = set(answer_data)

        if correct_answers == user_answers:
            is_correct = True
            points_earned = question.points

    elif question.question_type == 'matching':
        matches_data = json.loads(question.matches)

        # Проверяем новую структуру данных с поддержкой 3 колонок
        if 'correct_matches' in matches_data:
            correct_matches = matches_data['correct_matches']
            left_items = matches_data['left_items']
            columns = matches_data.get('columns', 2)
        else:
            # Старый формат (для обратной совместимости)
            correct_matches = {match['left']: {'right': match['right']} for match in matches_data}
            left_items = [match['left'] for match in matches_data]
            columns = 2

        user_correct_count = 0

        # Преобразуем пользовательские ответы в удобный формат
        user_matches = {}
        for pair in answer_data:
            left = pair['left']
            user_matches[left] = {}
            if 'middle' in pair:
                user_matches[left]['middle'] = pair['middle']
            if 'right' in pair:
                user_matches[left]['right'] = pair['right']

        # Проверяем правильность сопоставлений
        for left_item in left_items:
            if left_item in user_matches and left_item in correct_matches:
                correct_match = correct_matches[left_item]
                user_match = user_matches[left_item]

                if columns == 3:
                    # Трехколоночное сопоставление - должны совпадать обе колонки
                    if (user_match.get('middle') == correct_match.get('middle') and
                            user_match.get('right') == correct_match.get('right')):
                        user_correct_count += 1
                else:
                    # Двухколоночное сопоставление
                    if user_match.get('right') == correct_match.get('right'):
                        user_correct_count += 1

        # Если все левые элементы правильно сопоставлены
        if user_correct_count == len(left_items):
            is_correct = True
            points_earned = question.points
        else:
            # Частичные баллы за частично правильные ответы
            points_earned = (user_correct_count / len(left_items)) * question.points

    # Проверяем, есть ли уже ответ на этот вопрос
    existing_answer = Answer.query.filter_by(
        participation_id=participation.id,
        question_id=question_id
    ).first()

    if existing_answer:
        # Обновляем существующий ответ
        existing_answer.answer_data = json.dumps(answer_data)
        existing_answer.is_correct = is_correct
        existing_answer.points_earned = points_earned
        existing_answer.answered_at = get_current_time()
    else:
        # Создаем новый ответ
        answer = Answer(
            participation_id=participation.id,
            question_id=question_id,
            answer_data=json.dumps(answer_data),
            is_correct=is_correct,
            points_earned=points_earned
        )
        db.session.add(answer)

    # ВАЖНО: НЕ обновляем participation.total_points здесь!
    # Баллы будут подсчитаны в submit_block

    db.session.commit()

    return jsonify({'success': True, 'points': points_earned})


@app.route('/olympiad/<int:olympiad_id>/ranking', methods=['GET'])
@login_required
def get_ranking(olympiad_id):
    olympiad = Olympiad.query.get_or_404(olympiad_id)

    # Проверяем участие пользователя
    participation = Participation.query.filter_by(
        user_id=current_user.id,
        olympiad_id=olympiad_id
    ).first()

    if not participation:
        return jsonify({'success': False, 'message': 'Вы не участвуете в этой олимпиаде'})

    # Получаем текущий блок
    current_block = Block.query.get(participation.current_block)
    if not current_block:
        return jsonify({'success': False, 'message': 'Ошибка: блок не найден'})

    # Для незавершенных участий показываем статистику по текущему блоку
    if participation.status == 'in_progress':
        # Получаем баллы за текущий блок
        questions = Question.query.filter_by(block_id=current_block.id).all()
        answers = Answer.query.filter(
            Answer.participation_id == participation.id,
            Answer.question_id.in_([q.id for q in questions])
        ).all()

        block_points = sum(answer.points_earned for answer in answers)
        block_max_points = current_block.max_points

        # Для незавершенных участий не показываем место в общем рейтинге
        # Показываем только прогресс по текущему блоку
        response_data = {
            'success': True,
            'rank_position': 0,  # Не показываем место для незавершенных
            'rank_percentage': 0,
            'block_points': round(block_points, 1),
            'block_max_points': round(block_max_points, 1),
            'total_points': round(participation.total_points, 1),
            'total_participants': 1,
            'in_progress': True,
            'block_name': current_block.title
        }

        return jsonify(response_data)

    # Для завершенных участий показываем полную статистику
    # Обновляем итоговые баллы перед показом рейтинга
    update_all_final_scores(olympiad_id)

    # Получаем все завершенные участия для расчета места
    completed_participations = Participation.query.filter(
        Participation.olympiad_id == olympiad_id,
        Participation.status == 'completed'
    ).order_by(Participation.final_score.desc()).all()

    # Находим место текущего пользователя среди завершенных
    user_rank = 0
    for i, p in enumerate(completed_participations, 1):
        if p.id == participation.id:
            user_rank = i
            break

    # Вычисляем процент от максимально возможного места
    rank_percentage = 0
    if len(completed_participations) > 0 and user_rank > 0:
        rank_percentage = 100 - ((user_rank - 1) / len(completed_participations) * 100)

    # Получаем статистику по последнему завершенному блоку
    last_completed_block = None
    max_order = 0

    # Найдем блок с максимальным order, который пользователь завершил
    for block in Block.query.filter_by(olympiad_id=olympiad_id).order_by(Block.order).all():
        block_result = BlockResult.query.filter_by(
            participation_id=participation.id,
            block_id=block.id
        ).first()

        if block_result:
            last_completed_block = block
            max_order = block.order
        else:
            break

    if not last_completed_block:
        last_completed_block = Block.query.filter_by(olympiad_id=olympiad_id, order=1).first()

    # Баллы за последний завершенный блок
    block_result = BlockResult.query.filter_by(
        participation_id=participation.id,
        block_id=last_completed_block.id
    ).first()

    block_points = block_result.points_earned if block_result else 0

    response_data = {
        'success': True,
        'rank_position': user_rank,
        'rank_percentage': round(rank_percentage, 1),
        'block_points': round(block_points, 1),
        'block_max_points': round(last_completed_block.max_points, 1),
        'total_points': round(participation.final_score, 1),
        'total_participants': len(completed_participations),
        'in_progress': False,
        'block_name': last_completed_block.title
    }

    return jsonify(response_data)

# Исправленная функция submit_block - считаем баллы только здесь
@app.route('/olympiad/<int:olympiad_id>/submit_block', methods=['POST'])
@login_required
def submit_block(olympiad_id):
    # Проверяем участие пользователя
    participation = Participation.query.filter_by(
        user_id=current_user.id,
        olympiad_id=olympiad_id,
        status='in_progress'
    ).first()

    if not participation:
        return jsonify({'success': False, 'message': 'Вы не участвуете в этой олимпиаде'}), 403

    current_block = Block.query.get(participation.current_block)
    if not current_block:
        return jsonify({'success': False, 'message': 'Текущий блок не найден'}), 404

    # Проверяем, ответил ли пользователь на все вопросы блока
    questions = Question.query.filter_by(block_id=current_block.id).all()
    answered_questions = Answer.query.filter(
        Answer.participation_id == participation.id,
        Answer.question_id.in_([q.id for q in questions])
    ).count()

    if answered_questions < len(questions):
        return jsonify({
            'success': False,
            'message': f'Вы ответили только на {answered_questions} из {len(questions)} вопросов'
        })

    # Подсчитываем баллы за блок
    block_answers = Answer.query.filter(
        Answer.participation_id == participation.id,
        Answer.question_id.in_([q.id for q in questions])
    ).all()

    user_points = sum(answer.points_earned for answer in block_answers)
    total_points_possible = sum(q.points for q in questions)

    # ИСПРАВЛЕНО: правильно обновляем общий балл участника
    # Получаем старый результат блока
    old_block_result = BlockResult.query.filter_by(
        participation_id=participation.id,
        block_id=current_block.id
    ).first()

    if old_block_result:
        # Если блок уже был завершен ранее, вычитаем старые баллы
        participation.total_points -= old_block_result.points_earned
        old_block_result.points_earned = user_points
        old_block_result.completed_at = get_current_time()
    else:
        # Создаем новый результат блока
        block_result = BlockResult(
            participation_id=participation.id,
            block_id=current_block.id,
            points_earned=user_points,
            completed_at=get_current_time()
        )
        db.session.add(block_result)

    # Добавляем баллы за этот блок к общему счету
    participation.total_points += user_points

    # Проверяем процент правильных ответов
    percentage_correct = (user_points / total_points_possible) * 100 if total_points_possible > 0 else 0

    # Проверяем, достаточно ли баллов для перехода к следующему блоку
    if percentage_correct < current_block.threshold_percentage:
        # Недостаточно баллов, завершаем олимпиаду
        participation.status = 'completed'
        participation.finish_time = get_current_time()

        # Рассчитываем итоговый балл с учетом времени
        calculate_final_score(participation)

        db.session.commit()

        return jsonify({
            'success': True,
            'completed': True,
            'message': f'Вы набрали {percentage_correct:.1f}%, что меньше порогового значения {current_block.threshold_percentage}%. Олимпиада завершена.',
            'redirect': url_for('olympiad_results', olympiad_id=olympiad_id),
            'block_data': {
                'block_id': current_block.id,
                'block_name': current_block.title,
                'points_earned': round(user_points, 1),
                'total_points_possible': round(total_points_possible, 1),
                'percentage': round(percentage_correct, 1)
            }
        })

    # Ищем следующий блок
    next_block = Block.query.filter_by(
        olympiad_id=olympiad_id,
        order=current_block.order + 1
    ).first()

    if not next_block:
        # Это был последний блок, завершаем олимпиаду
        participation.status = 'completed'
        participation.finish_time = get_current_time()

        # Рассчитываем итоговый балл с учетом времени
        calculate_final_score(participation)

        db.session.commit()

        return jsonify({
            'success': True,
            'completed': True,
            'message': 'Поздравляем! Вы успешно завершили все блоки олимпиады.',
            'redirect': url_for('olympiad_results', olympiad_id=olympiad_id),
            'block_data': {
                'block_id': current_block.id,
                'block_name': current_block.title,
                'points_earned': round(user_points, 1),
                'total_points_possible': round(total_points_possible, 1),
                'percentage': round(percentage_correct, 1)
            }
        })

    # Переходим к следующему блоку
    participation.current_block = next_block.id

    # Обязательно фиксируем изменения в базе данных перед ответом
    db.session.commit()

    return jsonify({
        'success': True,
        'completed': False,
        'message': f'Вы успешно завершили блок "{current_block.title}" и набрали {percentage_correct:.1f}%. Переходим к следующему блоку.',
        'redirect': url_for('take_olympiad', olympiad_id=olympiad_id),
        'block_data': {
            'block_id': current_block.id,
            'block_name': current_block.title,
            'points_earned': round(user_points, 1),
            'total_points_possible': round(total_points_possible, 1),
            'percentage': round(percentage_correct, 1)
        }
    })



@app.route('/olympiad/<int:olympiad_id>/results', methods=['GET'])
@login_required
def olympiad_results(olympiad_id):
    olympiad = Olympiad.query.get_or_404(olympiad_id)

    # Получаем участие пользователя
    participation = Participation.query.filter_by(
        user_id=current_user.id,
        olympiad_id=olympiad_id
    ).first()

    if not participation or participation.status != 'completed':
        flash('Вы еще не завершили эту олимпиаду', 'error')
        return redirect(url_for('view_olympiad', olympiad_id=olympiad_id))

    # Обновляем итоговые баллы для всех участников
    update_all_final_scores(olympiad_id)

    # Получаем рейтинг на основе итогового балла
    rankings = Participation.query.filter_by(
        olympiad_id=olympiad_id,
        status='completed'
    ).order_by(Participation.final_score.desc()).all()

    user_rank = None
    for i, p in enumerate(rankings, 1):
        if p.id == participation.id:
            user_rank = i
            break

    # Детальная статистика по блокам
    blocks = Block.query.filter_by(olympiad_id=olympiad_id).order_by(Block.order).all()
    block_stats = []

    for block in blocks:
        questions = Question.query.filter_by(block_id=block.id).all()

        # Получаем ответы пользователя на вопросы этого блока
        answers = Answer.query.filter(
            Answer.participation_id == participation.id,
            Answer.question_id.in_([q.id for q in questions])
        ).all()

        # Если нет ответов на этот блок, значит пользователь до него не дошел
        if not answers:
            block_stats.append({
                'block': block,
                'attempted': False,
                'total_possible': sum(q.points for q in questions),
                'user_points': 0,
                'percentage': 0
            })
            continue

        total_possible = sum(q.points for q in questions)
        user_points = sum(a.points_earned for a in answers)
        percentage = (user_points / total_possible) * 100 if total_possible > 0 else 0

        block_stats.append({
            'block': block,
            'attempted': True,
            'total_possible': total_possible,
            'user_points': user_points,
            'percentage': percentage
        })

    return render_template(
        'olympiad/results.html',
        olympiad=olympiad,
        participation=participation,
        rankings=rankings,
        user_rank=user_rank,
        block_stats=block_stats,
        total_participants=len(rankings)
    )


@app.route('/admin/olympiad/<int:olympiad_id>/rankings', methods=['GET'])
@login_required
def admin_rankings(olympiad_id):
    if not current_user.is_admin:
        flash('У вас нет доступа к этой странице', 'error')
        return redirect(url_for('index'))

    olympiad = Olympiad.query.get_or_404(olympiad_id)

    # Обновляем итоговые баллы перед отображением
    update_all_final_scores(olympiad_id)

    # Получаем все завершенные участия, сортируем по итоговому баллу
    participations = Participation.query.filter_by(
        olympiad_id=olympiad_id,
        status='completed'
    ).order_by(Participation.final_score.desc()).all()

    # Получаем информацию о пользователях
    user_ids = [p.user_id for p in participations]
    users = {u.id: u for u in User.query.filter(User.id.in_(user_ids)).all()}

    return render_template(
        'admin/rankings.html',
        olympiad=olympiad,
        participations=participations,
        users=users
    )

@app.route('/olympiad/<int:olympiad_id>/finish_early', methods=['POST'])
@login_required
def finish_olympiad_early(olympiad_id):
    """Досрочное завершение олимпиады без временного бонуса"""
    # Проверяем участие пользователя
    participation = Participation.query.filter_by(
        user_id=current_user.id,
        olympiad_id=olympiad_id,
        status='in_progress'
    ).first()

    if not participation:
        return jsonify({'success': False, 'message': 'Вы не участвуете в этой олимпиаде'}), 403

    # Завершаем олимпиаду досрочно
    participation.status = 'completed'
    participation.finish_time = get_current_time()

    # При досрочном завершении временной бонус НЕ начисляется
    participation.time_bonus = 0
    participation.final_score = participation.total_points

    # Если есть duration_seconds, оставляем его для статистики
    if participation.start_time and participation.finish_time:
        duration = participation.finish_time - participation.start_time
        participation.duration_seconds = duration.total_seconds()

    db.session.commit()

    return jsonify({
        'success': True,
        'message': 'Олимпиада завершена досрочно. Временной бонус не начислен.',
        'redirect': url_for('olympiad_results', olympiad_id=olympiad_id)
    })


@app.route('/admin/olympiad/<int:olympiad_id>/export_pdf', methods=['GET'])
@login_required
def export_rankings_pdf(olympiad_id):
    if not current_user.is_admin:
        flash('У вас нет доступа к этой странице', 'error')
        return redirect(url_for('index'))

    olympiad = Olympiad.query.get_or_404(olympiad_id)

    # Обновляем итоговые баллы перед экспортом
    update_all_final_scores(olympiad_id)

    # Получаем все завершенные участия, сортируем по итоговому баллу
    participations = Participation.query.filter_by(
        olympiad_id=olympiad_id,
        status='completed'
    ).order_by(Participation.final_score.desc()).all()

    # Получаем информацию о пользователях
    user_ids = [p.user_id for p in participations]
    users = {u.id: u for u in User.query.filter(User.id.in_(user_ids)).all()}

    # Создаем HTML для PDF
    html = render_template(
        'admin/rankings_pdf.html',
        olympiad=olympiad,
        participations=participations,
        users=users
    )

    # Генерируем PDF
    pdf = pdfkit.from_string(html, False)

    # Отправляем PDF как файл
    buffer = BytesIO(pdf)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f'rankings_{olympiad.title}_{datetime.now().strftime("%Y%m%d")}.pdf',
        mimetype='application/pdf'
    )


# Добавляем роут для принудительного пересчета временных коэффициентов
@app.route('/admin/olympiad/<int:olympiad_id>/recalculate_scores', methods=['POST'])
@login_required
def recalculate_scores(olympiad_id):
    """Принудительный пересчет итоговых баллов с временным коэффициентом"""
    if not current_user.is_admin:
        return jsonify({'success': False, 'message': 'Доступ запрещен'}), 403

    try:
        update_all_final_scores(olympiad_id)
        return jsonify({'success': True, 'message': 'Итоговые баллы успешно пересчитаны'})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Ошибка при пересчете: {str(e)}'}), 500


if __name__ == '__main__':
    with app.app_context():
        db.create_all()

        try:
            # Попробуем выполнить запрос к новому столбцу
            db.session.execute('SELECT course FROM user LIMIT 1')
        except:
            # Если столбец не существует, добавляем его
            try:
                db.session.execute('ALTER TABLE user ADD COLUMN course INTEGER DEFAULT NULL')
                db.session.commit()
                print("Добавлен новый столбец 'course' в таблицу user")
            except:
                print("Столбец 'course' уже существует или произошла ошибка при добавлении")
        signatures_folder = 'static/signatures'
        if not os.path.exists(signatures_folder):
            os.makedirs(signatures_folder)
            print(f"Создана папка для подписей: {signatures_folder}")
            print("Поместите файлы подписей (1.jpg, 2.jpg, 3.jpg) в папку static/signatures/")
        # Проверяем и добавляем новые столбцы, если их нет
        try:
            # Попробуем выполнить запрос к новым столбцам
            db.session.execute('SELECT final_score, duration_seconds, time_bonus FROM participation LIMIT 1')
        except:
            # Если столбцы не существуют, добавляем их
            try:
                db.session.execute('ALTER TABLE participation ADD COLUMN final_score FLOAT DEFAULT 0')
                db.session.execute('ALTER TABLE participation ADD COLUMN duration_seconds FLOAT DEFAULT NULL')
                db.session.execute('ALTER TABLE participation ADD COLUMN time_bonus FLOAT DEFAULT 0')
                db.session.commit()
                print("Добавлены новые столбцы для временного коэффициента")
            except:
                print("Столбцы уже существуют или произошла ошибка при добавлении")

        # Создаем администратора, если его нет
        admin = User.query.filter_by(email='admin@example.com').first()
        if not admin:
            admin = User(
                email='admin@example.com',
                full_name='Administrator',
                is_admin=True
            )
            admin.set_password('admin')
            db.session.add(admin)
            db.session.commit()
            print("Создан администратор: admin@example.com / admin")

        # Пересчитываем итоговые баллы для всех существующих завершенных участий
        try:
            completed_participations = Participation.query.filter_by(status='completed').all()
            for participation in completed_participations:
                if participation.final_score == 0:  # Если еще не рассчитан
                    calculate_final_score(participation)
            db.session.commit()
            print(f"Пересчитаны итоговые баллы для {len(completed_participations)} участий")
        except Exception as e:
            print(f"Ошибка при пересчете существующих баллов: {e}")

    app.run(debug=True, host='0.0.0.0')